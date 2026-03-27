import logging
import json
import re
import requests
import pandas as pd
import io
import csv
from flask import Flask, session, request, render_template, send_file, send_from_directory, redirect, url_for
from datetime import datetime, timezone
from zoneinfo import ZoneInfo
from werkzeug.exceptions import HTTPException

# ----------------------------------------------------
# Comparison modules
# ----------------------------------------------------
from modules.group_rules import compare_group_rules
from modules.network_zones import compare_network_zones
from modules.session_policies import compare_session_policies
from modules.applications import compare_applications
from modules.authenticators import compare_authenticators
from modules.mfa_policies import compare_mfa_policies
from modules.password_policies import compare_password_policies
from modules.access_policies import compare_access_policies
from modules.idp_discovery_policies import compare_idp_discovery_policies
from modules.profile_enrollment_policies import compare_profile_enrollment_policies
from modules.brand_settings import compare_brand_settings
from modules.brand_pages import compare_brand_pages
from modules.brand_email_templates import compare_brand_email_templates
from modules.authorization_servers_settings import compare_authorization_servers_settings
from modules.authorization_servers_access_policies import compare_authorization_servers_access_policies
from modules.custom_admin_roles import compare_custom_admin_roles
from modules.resource_sets import compare_resource_sets
from modules.admin_assignments import compare_admin_assignments
from modules.api_tokens import compare_api_tokens
from modules.security_general_settings import compare_security_general_settings
from modules.org_settings import compare_org_settings
from modules.identity_providers import compare_identity_providers
from modules.realms import compare_realms, compare_realm_assignments
from modules.profile_schema_user import compare_user_profile_schema
from modules.profile_mappings import compare_profile_mappings
from modules.trusted_origins import compare_trusted_origins
from modules.event_hooks import compare_event_hooks
from modules.inline_hooks import compare_inline_hooks
from modules.attack_protection import compare_attack_protection
from modules.group_push_mappings import compare_group_push_mappings
from modules.entity_risk_policies import compare_entity_risk_policies
from modules.post_auth_session_policies import compare_post_auth_session_policies
from modules.agents import compare_agents
from modules.oktasnapshot_guide import build_oktasnapshot_guide

# ----------------------------------------------------
# Extractor modules
# ----------------------------------------------------
from scripts.extract_groups import get_groups
from scripts.extract_applications import get_applications as get_all_applications
from scripts.extract_users import get_users_with_security_context
from scripts.extract_api_tokens import get_api_tokens_with_metadata
from scripts.extract_admin_roles import (
    get_custom_admin_roles,
    get_resource_sets,
    get_resource_set_bindings,
)

app = Flask(__name__)
app.secret_key = "okta_compare_secret_key"
LAST_EXPORT = {"diffs": [], "matches": []}
OKTASNAPSHOT_EXPORT = {"rows": []}
OKTASNAPSHOT_GUIDE = {"sections": [], "domain": ""}
OKTAEVALUATE_EXPORT = {"evaluation": None}
OKTAMIGRATE_EXPORT = {
    "plan": None,
    "group_sync": None,
    "source_domain": "",
    "source_token": "",
    "target_domain": "",
    "target_token": "",
}

# ---------------------------------------------------
# Logging
# ---------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s - %(message)s",
)
logger = logging.getLogger("okta_compare")


def _section_map(sections):
    return {section.get("id"): section for section in (sections or [])}


def _ensure_https_domain(domain):
    return domain if str(domain).startswith(("http://", "https://")) else f"https://{domain}"


def _as_dict(value):
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, dict) else {}
        except Exception:
            return {}
    return {}


def _walk_nested(value, path=""):
    if isinstance(value, dict):
        for key, nested in value.items():
            next_path = f"{path}.{key}" if path else str(key)
            yield from _walk_nested(nested, next_path)
    elif isinstance(value, list):
        for idx, nested in enumerate(value):
            next_path = f"{path}[{idx}]"
            yield from _walk_nested(nested, next_path)
    else:
        yield path, value


def _is_deny_action(actions):
    for path, value in _walk_nested(actions):
        if isinstance(value, str) and value.strip().lower() == "deny":
            return True, path
    for path, value in _walk_nested(actions):
        if isinstance(value, str) and "deny" in value.strip().lower():
            return True, path
    return False, ""


def _extract_session_timeout_findings(actions, threshold_minutes=120):
    findings = []
    for path, value in _walk_nested(actions):
        if not isinstance(value, (int, float)):
            continue
        p = path.lower()
        if "session" not in p and "idle" not in p and "lifetime" not in p:
            continue

        minutes = None
        unit = None
        if "minute" in p:
            minutes = float(value)
            unit = "minutes"
        elif "second" in p:
            minutes = float(value) / 60.0
            unit = "seconds"
        elif "hour" in p:
            minutes = float(value) * 60.0
            unit = "hours"
        elif any(k in p for k in ["maxsession", "idle", "lifetime"]):
            # Fallback heuristic: assume minutes for session timeout-like fields without explicit unit
            minutes = float(value)
            unit = "assumed-minutes"

        if minutes is not None and minutes > threshold_minutes:
            findings.append(
                {
                    "path": path,
                    "value": value,
                    "unit": unit,
                    "minutes": round(minutes, 2),
                }
            )
    return findings


def _find_setting_value(rows, setting_name):
    for row in rows or []:
        if (row.get("Setting") or "").strip().lower() == setting_name.strip().lower():
            return row.get("Value")
    return None


def _status_from_boolish_enabled(value):
    if value is None:
        return "Info"
    normalized = str(value).strip().lower()
    if normalized in {"enabled", "yes", "true"}:
        return "Pass"
    if normalized in {"disabled", "not enabled", "no", "false"}:
        return "Fail"
    return "Info"


CHECK_PREFIX_LEGEND = [
    {"prefix": "ORG", "meaning": "Organization Settings"},
    {"prefix": "SEC", "meaning": "Security General Settings"},
    {"prefix": "GRP", "meaning": "Groups"},
    {"prefix": "GRR", "meaning": "Group Rules"},
    {"prefix": "NET", "meaning": "Network Zones"},
    {"prefix": "AUT", "meaning": "Authenticators"},
    {"prefix": "MFA", "meaning": "Authenticator Enrollment Policies"},
    {"prefix": "PWD", "meaning": "Password Policies"},
    {"prefix": "APP", "meaning": "Applications and App Sign-On"},
    {"prefix": "SES", "meaning": "Global Session Policies"},
    {"prefix": "POL", "meaning": "Cross-Policy and Risk Policy Checks"},
    {"prefix": "IDP", "meaning": "Identity Providers and IdP Discovery"},
    {"prefix": "AS", "meaning": "Authorization Servers"},
    {"prefix": "CAR", "meaning": "Custom Admin Roles"},
    {"prefix": "RST", "meaning": "Resource Sets"},
    {"prefix": "ADM", "meaning": "Admin Accounts and Admin Governance"},
    {"prefix": "GAD", "meaning": "Global Admin Accounts"},
    {"prefix": "API", "meaning": "API Tokens"},
    {"prefix": "BRD", "meaning": "Brand Settings"},
    {"prefix": "TOR", "meaning": "Trusted Origins"},
    {"prefix": "EVH", "meaning": "Event Hooks"},
    {"prefix": "INH", "meaning": "Inline Hooks"},
    {"prefix": "ATP", "meaning": "Attack Protection"},
    {"prefix": "RLM", "meaning": "Realms"},
    {"prefix": "PRF", "meaning": "Profile Schema and Mappings"},
    {"prefix": "GPM", "meaning": "Group Push Mappings"},
    {"prefix": "USR", "meaning": "User and Service Account Risk"},
]


CHECK_ID_BY_TITLE = {
    "Organization Support Metadata Completeness": "ORG-01",
    "Password Changed Notifications": "SEC-01",
    "Suspicious Activity Reporting for End Users": "SEC-02",
    "New Sign-On Notifications": "SEC-03",
    "Factor Enrollment Notifications": "SEC-04",
    "Factor Reset Notifications": "SEC-05",
    "ThreatInsight Blocking": "SEC-06",
    "Groups Missing Description": "GRP-01",
    "Disabled Group Rules": "GRR-01",
    "Blocklisted Network Zone Presence": "NET-01",
    "Network Zones Defined But Not Used In Policies": "NET-02",
    "Trusted Network Zones Without Entries": "NET-03",
    "Weak Authenticators Enabled": "AUT-01",
    "Phishing-Resistant Authenticator Availability": "AUT-02",
    "Weaker Factors in MFA Enrollment Policies": "MFA-01",
    "Optional Factors in MFA Enrollment Policies": "MFA-02",
    "Required Authenticator in MFA Enrollment Policies": "MFA-03",
    "Password Policy Strength": "PWD-01",
    "SAML Authentication Supported but Disabled Apps": "APP-01",
    "Applications Assigned to Everyone": "APP-02",
    "Password-Based Application Sign-On Modes Present": "APP-03",
    "App Sign-On Policy Catch-All Deny": "APP-04",
    "Session Lifetime <= 2 Hours": "SES-01",
    "High-Risk Request MFA Every Sign-In": "POL-01",
    "New Device MFA Every Sign-In": "POL-02",
    "Admin Console MFA Every Sign-In": "POL-03",
    "Entity Risk Policy Rule Coverage": "POL-04",
    "Identity Threat Protection Policy Rule Coverage": "POL-05",
    "Inactive Identity Providers": "IDP-01",
    "Inactive IdP Discovery Rules": "IDP-02",
    "Authorization Servers Without Automatic Key Rotation": "AS-01",
    "Authorization Server Access Rules With Broad Client Scope": "AS-02",
    "Custom Admin Roles Missing Description": "CAR-01",
    "Resource Sets Without Resources": "RST-01",
    "Resource Sets Without Bindings": "RST-02",
    "Admin Public Client Applications Present": "ADM-01",
    "No MFA - Admin Account": "ADM-02",
    "Pending MFA - Admin Account": "ADM-03",
    "No MFA Enforced - Admin Account": "ADM-04",
    "SSO Bypass - Admin": "ADM-05",
    "SSO Bypass + No MFA - Admin": "ADM-06",
    "Old Password - Admin Account": "ADM-07",
    "Old Password - Admin Service Account": "ADM-08",
    "Unused Admin Account": "ADM-09",
    "Unused Admin Service Account": "ADM-10",
    "Unused Administrative Roles": "ADM-11",
    "Service Account with Console Access - Admin": "ADM-12",
    "Old Password, No MFA, Unused Admin Accounts": "ADM-13",
    "Old Password, No MFA, Unused Admin Service Accounts": "ADM-14",
    "Unrotated and Unused Keys and Tokens - Admin": "ADM-15",
    "Unrotated Keys and Tokens - Admin": "ADM-16",
    "Unused Keys and Tokens - Admins": "ADM-17",
    "No MFA - Global Admin Account": "GAD-01",
    "Pending MFA - Global Admin Account": "GAD-02",
    "No MFA Enforced - Global Admin Account": "GAD-03",
    "Old Password - Global Admin Account": "GAD-04",
    "Old Password - Global Admin Service Account": "GAD-05",
    "Unused Global Admin Account": "GAD-06",
    "Unused Global Admin Service Account": "GAD-07",
    "Excessive Number of Super Admins": "GAD-08",
    "Super Admin with API Token": "GAD-09",
    "API Tokens Without Network Restrictions": "API-01",
    "Brands Missing Custom Privacy Policy URL": "BRD-01",
    "Insecure Trusted Origins": "TOR-01",
    "Unverified Event Hooks": "EVH-01",
    "Event Hooks Without Authentication Scheme": "EVH-02",
    "Inactive Event Hooks": "EVH-03",
    "Inline Hooks Without Authentication Scheme": "INH-01",
    "Inactive Inline Hooks": "INH-02",
    "Attack Protection Controls Not Enforcing": "ATP-01",
    "Inactive Realms": "RLM-01",
    "Realm Default Assignment Coverage": "RLM-02",
    "Sensitive Writable Profile Attributes": "PRF-01",
    "Empty Profile Mappings": "PRF-02",
    "Inactive Group Push Mappings": "GPM-01",
    "Stale Group Push Mappings": "GPM-02",
    "No MFA - Account": "USR-01",
    "Pending MFA": "USR-02",
    "SSO Bypass - Account": "USR-03",
    "SSO Bypass + No MFA - Account": "USR-04",
    "Old Password - Account": "USR-05",
    "Old Password - Service Account": "USR-06",
    "Unused Account": "USR-07",
    "Unused Service Account": "USR-08",
    "Partially Off-boarded User": "USR-09",
    "Service Account with Console Access - Account": "USR-10",
    "Old Password, No MFA, Unused Accounts": "USR-11",
    "Old Password, No MFA, Unused Service Accounts": "USR-12",
    "Unrotated and Unused Keys and Tokens - Account": "USR-13",
    "Unrotated Keys and Tokens - Account": "USR-14",
    "Unused Keys and Tokens - Account": "USR-15",
}


def _validation(title, status, summary, items=None, severity=None):
    check_id = CHECK_ID_BY_TITLE.get(title, "")
    return {
        "check_id": check_id,
        "title": title,
        "status": status,
        "severity": severity,
        "summary": summary,
        "items": items or [],
    }


def _extract_minutes_from_label(text):
    if text is None:
        return None
    s = str(text).strip().lower()
    if not s:
        return None
    m = re.search(r"(\d+)\s*hour", s)
    if m:
        return int(m.group(1)) * 60
    m = re.search(r"(\d+)\s*minute", s)
    if m:
        return int(m.group(1))
    return None


def _password_policy_weaknesses(policy_row):
    weaknesses = []
    complexity = str(policy_row.get("Complexity Settings") or "")
    age = str(policy_row.get("Age Settings") or "")
    lockout = str(policy_row.get("Lockout Settings") or "")

    min_len = None
    m = re.search(r"Minimum length:\s*(\d+)", complexity, re.IGNORECASE)
    if m:
        min_len = int(m.group(1))
        if min_len < 12:
            weaknesses.append(f"minimum length {min_len} (< 12)")

    for label in ["Lower case letters", "Upper case letters", "Numbers", "Symbols"]:
        m = re.search(rf"{re.escape(label)}:\s*(\d+)", complexity, re.IGNORECASE)
        if m and int(m.group(1)) == 0:
            weaknesses.append(f"{label.lower()} requirement is 0")

    if "Restrict use of common passwords: Disabled" in complexity:
        weaknesses.append("common password restriction disabled")
    if "Does not contain part of username: No" in complexity:
        weaknesses.append("username exclusion disabled")

    min_age = _extract_minutes_from_label(re.search(r"Minimum password age:\s*([^;]+)", age, re.IGNORECASE).group(1)) if re.search(r"Minimum password age:\s*([^;]+)", age, re.IGNORECASE) else None
    if min_age is not None and min_age == 0:
        weaknesses.append("minimum password age is 0")

    history_count = re.search(r"Enforce password history \(count\):\s*(\d+)", age, re.IGNORECASE)
    if history_count and int(history_count.group(1)) < 24:
        weaknesses.append(f"password history {history_count.group(1)} (< 24)")

    lockout_attempts = re.search(r"Lock out after failed attempts:\s*(\d+)", lockout, re.IGNORECASE)
    if lockout_attempts and int(lockout_attempts.group(1)) > 10:
        weaknesses.append(f"lockout attempts threshold {lockout_attempts.group(1)} (> 10)")

    return weaknesses


def _mfa_policy_has_optional_factors(policy_row):
    settings = _as_dict(policy_row.get("Settings"))
    for _, value in _walk_nested(settings):
        if isinstance(value, str) and value.strip().upper() == "OPTIONAL":
            return True
    return False


def _mfa_policy_weak_factors(policy_row):
    settings = _as_dict(policy_row.get("Settings"))
    weak_factor_hits = set()
    weak_tokens = {
        "sms": "SMS",
        "voice": "Voice",
        "call": "Voice",
        "security_question": "Security Question",
        "question": "Security Question",
        "email": "Email",
    }
    for path, value in _walk_nested(settings):
        p = path.lower()
        matched = None
        for token, label in weak_tokens.items():
            if token in p:
                matched = label
                break
        if not matched:
            continue
        if isinstance(value, str):
            v = value.strip().upper()
            if v in {"OPTIONAL", "REQUIRED", "ACTIVE", "ENABLED"}:
                weak_factor_hits.add(matched)
        elif isinstance(value, bool) and value:
            weak_factor_hits.add(matched)
    return sorted(weak_factor_hits)


def _section_rows(section_by_id, section_id, entry_type=None):
    rows = (section_by_id.get(section_id) or {}).get("rows") or []
    if entry_type is None:
        return rows
    return [row for row in rows if str(row.get("Entry Type") or "").strip() == entry_type]


def _row_map(rows, key_field="Setting", value_field="Value"):
    mapping = {}
    for row in rows or []:
        key = row.get(key_field)
        if key:
            mapping[str(key)] = row.get(value_field)
    return mapping


def _blankish(value):
    if value is None:
        return True
    if isinstance(value, str):
        return value.strip() == "" or value.strip().lower() in {"not available", "none", "null"}
    if isinstance(value, (list, dict, tuple, set)):
        return len(value) == 0
    return False


def _text(value):
    if isinstance(value, (dict, list)):
        return json.dumps(value, sort_keys=True, default=str)
    return str(value or "")


def _contains_token(value, tokens):
    haystack = _text(value).lower()
    return any(token.lower() in haystack for token in tokens)


def _looks_like_inactive(status):
    return str(status or "").strip().upper() not in {"ACTIVE", "ENABLED", "VERIFIED"}


def _parse_iso_datetime(value):
    if value in (None, ""):
        return None
    text = str(value).strip()
    if not text:
        return None
    if text.endswith("Z"):
        text = f"{text[:-1]}+00:00"
    try:
        parsed = datetime.fromisoformat(text)
    except ValueError:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed


def _days_since(value):
    parsed = _parse_iso_datetime(value)
    if not parsed:
        return None
    return (datetime.now(timezone.utc) - parsed.astimezone(timezone.utc)).days


def _user_display_name(user):
    profile = user.get("profile") or {}
    display_name = (profile.get("displayName") or "").strip()
    if display_name:
        return display_name
    full_name = " ".join(
        part.strip()
        for part in [profile.get("firstName") or "", profile.get("lastName") or ""]
        if str(part).strip()
    ).strip()
    if full_name:
        return full_name
    return str(profile.get("login") or user.get("id") or "Unknown User")


def _user_identifier(user):
    profile = user.get("profile") or {}
    login = profile.get("login") or profile.get("email") or user.get("id") or "Unknown User"
    display_name = _user_display_name(user)
    if display_name and display_name != login:
        return f"{login} ({display_name})"
    return str(login)


def _user_provider_type(user):
    credentials = user.get("credentials") or {}
    provider = credentials.get("provider") or {}
    return str(provider.get("type") or "").strip().upper()


def _user_role_types(user):
    role_types = set()
    for role in user.get("roles") or []:
        role_type = str(role.get("type") or role.get("label") or role.get("name") or "").strip().upper()
        if role_type:
            role_types.add(role_type)
    return role_types


def _active_factor_types(user):
    factor_types = set()
    for factor in user.get("factors") or []:
        status = str(factor.get("status") or "").strip().upper()
        if status != "ACTIVE":
            continue
        factor_type = str(factor.get("factorType") or factor.get("factor") or "").strip()
        provider = str(factor.get("provider") or "").strip()
        label = f"{factor_type} ({provider})" if provider else factor_type
        if label:
            factor_types.add(label)
    return factor_types


def _pending_factor_types(user):
    factor_types = set()
    for factor in user.get("factors") or []:
        status = str(factor.get("status") or "").strip().upper()
        if status not in {"PENDING", "PENDING_ACTIVATION"}:
            continue
        factor_type = str(factor.get("factorType") or factor.get("factor") or "").strip()
        provider = str(factor.get("provider") or "").strip()
        label = f"{factor_type} ({provider})" if provider else factor_type
        if label:
            factor_types.add(label)
    return factor_types


def _has_mfa(user):
    return bool(_active_factor_types(user))


def _has_pending_mfa(user):
    return (not _has_mfa(user)) and bool(_pending_factor_types(user))


def _is_direct_access_user(user):
    provider_type = _user_provider_type(user)
    if provider_type in {"FEDERATION", "SOCIAL"}:
        return False
    if provider_type:
        return True
    return bool((user.get("credentials") or {}).get("password"))


def _is_service_account(user, active_token_count=0):
    profile = user.get("profile") or {}
    searchable = " ".join(
        str(value or "")
        for value in [
            profile.get("login"),
            profile.get("email"),
            profile.get("displayName"),
            profile.get("firstName"),
            profile.get("lastName"),
            profile.get("title"),
            profile.get("department"),
        ]
    ).lower()
    patterns = [
        r"(^|[\W_])(svc|service|bot|automation|api|daemon|integration|system|noreply|nonhuman|robot)([\W_]|$)",
        r"(^|[\W_])(ci|cd|etl|job)([\W_]|$)",
    ]
    if any(re.search(pattern, searchable) for pattern in patterns):
        return True
    if active_token_count > 0 and not str(profile.get("firstName") or "").strip() and not str(profile.get("lastName") or "").strip():
        return True
    return False


def _is_unused_user(user, threshold_days=91):
    status = str(user.get("status") or "").strip().upper()
    if status not in {"ACTIVE", "PROVISIONED", "RECOVERY", "LOCKED_OUT", "PASSWORD_EXPIRED", "STAGED"}:
        return False
    days_since_last_login = _days_since(user.get("lastLogin"))
    if days_since_last_login is not None:
        return days_since_last_login >= threshold_days
    created_or_activated_days = _days_since(user.get("activated")) or _days_since(user.get("created"))
    return created_or_activated_days is not None and created_or_activated_days >= threshold_days


def _has_old_password(user, threshold_days=90):
    if _user_provider_type(user) in {"FEDERATION", "SOCIAL"}:
        return False
    age_days = _days_since(user.get("passwordChanged"))
    return age_days is not None and age_days >= threshold_days


def _token_rotation_age_days(token):
    for candidate in [token.get("lastUpdated"), token.get("created")]:
        days = _days_since(candidate)
        if days is not None:
            return days
    return None


def _is_active_token(token):
    return str(token.get("status") or "").strip().upper() == "ACTIVE"


def _token_name(token):
    return str(token.get("name") or token.get("label") or token.get("id") or "Unknown Token")


def _token_inventory(tokens):
    active_by_user = {}
    for token in tokens or []:
        if not _is_active_token(token):
            continue
        user_id = token.get("userId")
        if not user_id:
            continue
        active_by_user.setdefault(user_id, []).append(token)
    return active_by_user


def _token_findings_for_user(user_tokens, owner_unused=False, rotation_days=90):
    unrotated = []
    unused = []
    for token in user_tokens or []:
        age_days = _token_rotation_age_days(token)
        if age_days is not None and age_days >= rotation_days:
            unrotated.append((token, age_days))
        if owner_unused:
            unused.append(token)
    return unrotated, unused


def _format_token_item(user, token, suffix=None):
    base = f"{_user_identifier(user)}: {_token_name(token)}"
    if suffix:
        return f"{base} ({suffix})"
    return base


def _role_binding_keys(binding_role):
    keys = set()
    if isinstance(binding_role, dict):
        for value in [binding_role.get("id"), binding_role.get("type"), binding_role.get("label"), binding_role.get("name")]:
            if value:
                keys.add(str(value).strip())
                keys.add(str(value).strip().upper())
    elif binding_role:
        keys.add(str(binding_role).strip())
        keys.add(str(binding_role).strip().upper())
    return keys


def _custom_role_usage_map(custom_roles, resource_set_bindings):
    usage = {}
    for role in custom_roles or []:
        keys = set()
        for value in [role.get("id"), role.get("label"), role.get("name")]:
            if value:
                keys.add(str(value).strip())
                keys.add(str(value).strip().upper())
        usage[str(role.get("id") or role.get("label") or role.get("name") or "")] = {
            "role": role,
            "keys": keys,
            "used": False,
        }

    for binding in resource_set_bindings or []:
        binding_keys = _role_binding_keys(binding.get("role"))
        if not binding_keys:
            continue
        for role_usage in usage.values():
            if role_usage["keys"] & binding_keys:
                role_usage["used"] = True

    return usage


def _actions_reference_mfa(actions):
    positive_strings = {
        "mfa",
        "multifactor",
        "promptforfactor",
        "phishing_resistant",
        "authenticator",
        "factor",
        "password / idp + another factor",
    }
    positive_paths = {"mfa", "factor", "authenticator", "challenge"}
    for path, value in _walk_nested(actions):
        path_l = path.lower()
        value_l = str(value).strip().lower()
        if any(token in path_l for token in positive_paths):
            return True
        if any(token in value_l for token in positive_strings):
            return True
    return False


def _actions_reference_every_sign_in(actions):
    positive_values = {
        "every sign in",
        "every sign-in",
        "every sign-in attempt",
        "at every sign in",
        "at every sign-in",
        "always",
        "every_time",
        "everytime",
        "per_session",
    }
    for path, value in _walk_nested(actions):
        path_l = path.lower()
        value_l = str(value).strip().lower()
        if any(token in path_l for token in ["prompt", "reauth", "factorlifetime", "mfalifetime"]):
            if any(token in value_l for token in positive_values):
                return True
        if any(token in value_l for token in positive_values):
            return True
    return False


def _conditions_reference_new_device(conditions):
    return _contains_token(conditions, ["new device", "new_device", "behavior", "device"])


def _conditions_reference_high_risk(conditions):
    return _contains_token(conditions, ['"high"', "risk", "high"])


def _policy_has_required_authenticator(policy_row):
    settings = _as_dict(policy_row.get("Settings"))
    for _, value in _walk_nested(settings):
        if isinstance(value, str) and value.strip().upper() == "REQUIRED":
            return True
    return False


def _tenant_mfa_enforcement_gaps(section_by_id):
    mfa_rows = (section_by_id.get("mfa-enrollment-policies") or {}).get("rows") or []
    auth_rows = (section_by_id.get("authentication-policies") or {}).get("rows") or []
    mfa_policy_rows = [r for r in mfa_rows if r.get("Entry Type") == "Policy"]
    auth_rule_rows = [r for r in auth_rows if r.get("Entry Type") == "Rule"]

    optional_enrollment = any(_mfa_policy_has_optional_factors(row) for row in mfa_policy_rows)
    auth_requires_mfa = any(_actions_reference_mfa(_as_dict(row.get("Actions"))) for row in auth_rule_rows)

    return {
        "optional_enrollment": optional_enrollment,
        "auth_requires_mfa": auth_requires_mfa,
        "has_data": bool(mfa_policy_rows or auth_rule_rows),
    }


def _identity_validation(title, severity, records, warning_summary, pass_summary, empty_summary=None, data_available=True):
    if records:
        return _validation(title, "Fail", warning_summary(records), records[:20], severity=severity)
    if not data_available and empty_summary:
        return _validation(title, "Pass", empty_summary, severity=severity)
    return _validation(title, "Pass", pass_summary, severity=severity)


def _unsupported_ispm_validations():
    return []


def _run_security_validations(sections, extra_context=None):
    section_by_id = _section_map(sections)
    extra_context = extra_context or {}
    validations = []

    # Validation 1: Catch-all/default rule deny posture in authentication/app sign-on policies
    auth_section = section_by_id.get("authentication-policies") or {}
    auth_rows = auth_section.get("rows") or []
    auth_rule_rows = [r for r in auth_rows if (r.get("Entry Type") == "Rule")]

    policy_map = {}
    for row in auth_rule_rows:
        policy_name = row.get("Policy Name") or "Unknown Policy"
        rule_name = (row.get("Rule Name") or "").strip()
        if not rule_name:
            continue
        bucket = policy_map.setdefault(policy_name, {"catch_all_found": False, "catch_all_denied": False, "details": []})
        is_catch_all = any(token in rule_name.lower() for token in ["catch", "default", "all"])
        actions = _as_dict(row.get("Actions"))
        is_deny, deny_path = _is_deny_action(actions)
        if is_catch_all:
            bucket["catch_all_found"] = True
            if is_deny:
                bucket["catch_all_denied"] = True
            bucket["details"].append(
                {
                    "rule_name": rule_name,
                    "status": row.get("Status"),
                    "deny": is_deny,
                    "deny_path": deny_path,
                }
            )

    missing_or_not_deny = []
    passing = []
    for policy_name, result in policy_map.items():
        if result["catch_all_found"] and result["catch_all_denied"]:
            passing.append(policy_name)
        else:
            missing_or_not_deny.append(
                {
                    "policy": policy_name,
                    "issue": "Catch-all/default rule not found or not set to deny",
                    "details": result["details"],
                }
            )

    if not policy_map:
        validations.append(_validation(
            "App Sign-On Policy Catch-All Deny",
            "Pass",
            "No authentication/app sign-on policy rule data was available to validate catch-all deny posture.",
            severity="High",
        ))
    elif missing_or_not_deny:
        validations.append(_validation(
            "App Sign-On Policy Catch-All Deny",
            "Fail",
            f"{len(missing_or_not_deny)} policy(s) may not have a catch-all/default deny rule configured.",
            [f"{item['policy']}: {item['issue']}" for item in missing_or_not_deny],
            severity="High",
        ))
    else:
        validations.append(_validation(
            "App Sign-On Policy Catch-All Deny",
            "Pass",
            f"Validated catch-all/default deny posture for {len(passing)} authentication/app sign-on policy(ies).",
            [f"{name}: catch-all/default deny detected" for name in passing[:8]],
            severity="High",
        ))

    # Validation 2: Session timeout > 2 hours in global session policies
    session_section = section_by_id.get("global-session-policies") or {}
    session_rows = session_section.get("rows") or []
    session_rule_rows = [r for r in session_rows if r.get("Entry Type") == "Rule"]
    timeout_violations = []
    for row in session_rule_rows:
        actions = _as_dict(row.get("Actions"))
        for finding in _extract_session_timeout_findings(actions, threshold_minutes=120):
            timeout_violations.append(
                {
                    "policy": row.get("Policy Name") or "Unknown Policy",
                    "rule": row.get("Rule Name") or "Unnamed Rule",
                    "path": finding["path"],
                    "minutes": finding["minutes"],
                    "raw": f"{finding['value']} ({finding['unit']})",
                }
            )

    if not session_rule_rows:
        validations.append(_validation(
            "Session Lifetime <= 2 Hours",
            "Pass",
            "No global session policy rule data was available to validate session timeout settings.",
            severity="High",
        ))
    elif timeout_violations:
        policy_count = len({(v["policy"]) for v in timeout_violations})
        validations.append(_validation(
            "Session Lifetime <= 2 Hours",
            "Fail",
            f"Session lifetime duration is more than 2 hours on {policy_count} policy(s).",
            [
                f"{v['policy']} / {v['rule']}: {v['path']} = {v['raw']} (~{int(v['minutes'])} min)"
                for v in timeout_violations[:20]
            ],
            severity="High",
        ))
    else:
        validations.append(_validation(
            "Session Lifetime <= 2 Hours",
            "Pass",
            "No session timeout values above 2 hours were detected in global session policy rules.",
            severity="High",
        ))

    # Security notifications checks (High)
    sec_rows = (section_by_id.get("security-settings") or {}).get("rows") or []
    notification_checks = [
        ("Password Changed Notifications", "Password changed notification email"),
        ("Suspicious Activity Reporting for End Users", "Report suspicious activity via email"),
        ("New Sign-On Notifications", "New sign-on notification email"),
        ("Factor Enrollment Notifications", "Authenticator enrolled notification email"),
        ("Factor Reset Notifications", "Authenticator reset notification email"),
    ]
    for title, setting_name in notification_checks:
        if not sec_rows:
            validations.append(_validation(
                title,
                "Pass",
                f"No security settings data was available to validate {title.lower()}.",
                severity="High",
            ))
            continue
        value = _find_setting_value(sec_rows, setting_name)
        status = _status_from_boolish_enabled(value)
        if status == "Pass":
            summary = f"{title} are enabled."
        elif status == "Fail":
            summary = f"{title} are disabled."
        else:
            summary = f"{title} setting was not available in tenant response."
        validations.append(_validation(title, status, summary, severity="High"))

    threatinsight_action = _find_setting_value(sec_rows, "Action")
    if not sec_rows:
        validations.append(_validation(
            "ThreatInsight Blocking",
            "Pass",
            "No security settings data was available to validate ThreatInsight blocking.",
            severity="High",
        ))
    elif threatinsight_action is None:
        validations.append(_validation(
            "ThreatInsight Blocking",
            "Info",
            "ThreatInsight action setting was not available in tenant response.",
            severity="High",
        ))
    elif "enforce" in str(threatinsight_action).lower() or "block" in str(threatinsight_action).lower():
        validations.append(_validation(
            "ThreatInsight Blocking",
            "Pass",
            "ThreatInsight is configured to block or enforce security on suspicious IP activity.",
            severity="High",
        ))
    else:
        validations.append(_validation(
            "ThreatInsight Blocking",
            "Fail",
            "ThreatInsight is not configured in blocking/enforcement mode.",
            [str(threatinsight_action)],
            severity="High",
        ))

    # Weak password policies (Moderate)
    pwd_rows = (section_by_id.get("password-policies") or {}).get("rows") or []
    pwd_policy_rows = [r for r in pwd_rows if r.get("Entry Type") == "Policy"]
    weak_pwd = []
    for row in pwd_policy_rows:
        weaknesses = _password_policy_weaknesses(row)
        if weaknesses:
            weak_pwd.append((row.get("Name") or row.get("Policy Name") or "Unnamed Policy", weaknesses))
    if not pwd_policy_rows:
        validations.append(_validation(
            "Password Policy Strength",
            "Pass",
            "No password policy data was available for password strength validation.",
            severity="Moderate",
        ))
    elif weak_pwd:
        validations.append(_validation(
            "Password Policy Strength",
            "Fail",
            f"Password policies for {len(weak_pwd)} policy(s) are weak.",
            [f"{name}: {', '.join(issues[:3])}" for name, issues in weak_pwd],
            severity="Moderate",
        ))
    else:
        validations.append(_validation(
            "Password Policy Strength",
            "Pass",
            "No weak password policy patterns were detected using current heuristics.",
            severity="Moderate",
        ))

    # Network zones blocklist presence (Moderate)
    zone_rows = (section_by_id.get("network-zones") or {}).get("rows") or []
    has_block_zone = any(
        "block" in str((z.get("Usage") or "")).lower() or "block" in str((z.get("Name") or "")).lower()
        for z in zone_rows
    )
    if not zone_rows:
        validations.append(_validation(
            "Blocklisted Network Zone Presence",
            "Pass",
            "No network zone data was available to validate blocklisted zone configuration.",
            severity="Moderate",
        ))
    elif not has_block_zone:
        validations.append(_validation(
            "Blocklisted Network Zone Presence",
            "Fail",
            "Network Zones do not contain a block listed zone.",
            severity="Moderate",
        ))
    else:
        validations.append(_validation(
            "Blocklisted Network Zone Presence",
            "Pass",
            "At least one blocklisted network zone was detected.",
            severity="Moderate",
        ))

    # MFA enrollment policies: weaker factors + optional factors
    mfa_rows = (section_by_id.get("mfa-enrollment-policies") or {}).get("rows") or []
    mfa_policy_rows = [r for r in mfa_rows if r.get("Entry Type") == "Policy"]
    weak_factor_policies = []
    optional_factor_policies = []
    for row in mfa_policy_rows:
        name = row.get("Name") or row.get("Policy Name") or "Unnamed Policy"
        weak_factors = _mfa_policy_weak_factors(row)
        if weak_factors:
            weak_factor_policies.append((name, weak_factors))
        if _mfa_policy_has_optional_factors(row):
            optional_factor_policies.append(name)
    policies_without_required_factor = [
        row.get("Name") or row.get("Policy Name") or "Unnamed Policy"
        for row in mfa_policy_rows
        if not _policy_has_required_authenticator(row)
    ]

    if not mfa_policy_rows:
        validations.append(_validation(
            "Weaker Factors in MFA Enrollment Policies",
            "Pass",
            "No MFA enrollment policy data was available to evaluate weaker factor usage.",
            severity="High",
        ))
        validations.append(_validation(
            "Optional Factors in MFA Enrollment Policies",
            "Pass",
            "No MFA enrollment policy data was available to evaluate optional factor enrollment.",
            severity="Moderate",
        ))
        validations.append(_validation(
            "Required Authenticator in MFA Enrollment Policies",
            "Pass",
            "No MFA enrollment policy data was available to evaluate required authenticator posture.",
            severity="High",
        ))
    else:
        validations.append(_validation(
            "Weaker Factors in MFA Enrollment Policies",
            "Fail" if weak_factor_policies else "Pass",
            (f"Weaker factors are set in {len(weak_factor_policies)} policies."
             if weak_factor_policies else "No weaker factors detected in MFA enrollment policy settings."),
            [f"{name}: {', '.join(factors)}" for name, factors in weak_factor_policies[:20]],
            severity="High",
        ))
        validations.append(_validation(
            "Optional Factors in MFA Enrollment Policies",
            "Fail" if optional_factor_policies else "Pass",
            (f"Factors are optional for {len(optional_factor_policies)} Factor Enrollment policies."
             if optional_factor_policies else "No optional factor enrollment settings detected in MFA enrollment policies."),
            optional_factor_policies[:20],
            severity="Moderate",
        ))
        validations.append(_validation(
            "Required Authenticator in MFA Enrollment Policies",
            "Fail" if policies_without_required_factor else "Pass",
            (
                f"{len(policies_without_required_factor)} MFA enrollment policy(ies) do not appear to require any authenticator."
                if policies_without_required_factor
                else "Every MFA enrollment policy appears to require at least one authenticator."
            ),
            policies_without_required_factor[:20],
            severity="High",
        ))

    # SAML apps disabled (High) using full app inventory (including inactive)
    all_apps = extra_context.get("all_apps") or []
    if not all_apps:
        validations.append(_validation(
            "SAML Authentication Supported but Disabled Apps",
            "Pass",
            "No full application inventory was available to validate disabled SAML apps.",
            severity="High",
        ))
    else:
        disabled_saml = [
            app for app in all_apps
            if str(app.get("signOnMode") or "").upper() == "SAML_2_0"
            and str(app.get("status") or "").upper() != "ACTIVE"
        ]
        validations.append(_validation(
            "SAML Authentication Supported but Disabled Apps",
            "Fail" if disabled_saml else "Pass",
            (f"SAML authentication is supported but disabled for {len(disabled_saml)} apps."
             if disabled_saml else "No disabled SAML applications were detected."),
            [str(app.get("label") or app.get("name") or app.get("id")) for app in disabled_saml[:20]],
            severity="High",
        ))

    # Org settings readiness checks
    org_setting_map = _row_map(_section_rows(section_by_id, "org-settings"))
    org_missing = [
        label for label in [
            "Company Name",
            "Website",
            "End User Support Help URL",
            "Billing Contact Email",
            "Technical Contact Email",
        ]
        if _blankish(org_setting_map.get(label))
    ]
    validations.append(_identity_validation(
        "Organization Support Metadata Completeness",
        "Moderate",
        org_missing,
        lambda items: f"{len(items)} important org support/contact setting(s) are missing.",
        "Core org support and contact metadata are populated.",
        "Organization settings were not available for validation.",
        data_available=bool(_section_rows(section_by_id, "org-settings")),
    ))

    # Group hygiene checks
    group_rows = _section_rows(section_by_id, "groups")
    groups_missing_description = [
        str(row.get("Group Name") or "Unnamed Group")
        for row in group_rows
        if _blankish(row.get("Description"))
    ]
    validations.extend([
        _identity_validation(
            "Groups Missing Description",
            "Low",
            groups_missing_description,
            lambda items: f"{len(items)} group(s) do not have a description.",
            "All Okta groups have descriptions.",
            "No group data was available for group description validation.",
            data_available=bool(group_rows),
        ),
    ])

    # Group rule hygiene
    group_rule_rows = _section_rows(section_by_id, "group-rules")
    disabled_group_rules = [
        str(row.get("Rule Name") or "Unnamed Rule")
        for row in group_rule_rows
        if _looks_like_inactive(row.get("Status"))
    ]
    validations.extend([
        _identity_validation(
            "Disabled Group Rules",
            "Low",
            disabled_group_rules,
            lambda items: f"{len(items)} group rule(s) are not active.",
            "All group rules are active.",
            "No group rule data was available for rule status validation.",
            data_available=bool(group_rule_rows),
        ),
    ])

    # Network zones
    zone_rows = _section_rows(section_by_id, "network-zones")
    network_policy_rows = _section_rows(section_by_id, "global-session-policies", entry_type="Rule") + _section_rows(
        section_by_id,
        "authentication-policies",
        entry_type="Rule",
    )
    zone_usage_blob = " ".join(
        json.dumps(row.get("Conditions Network"), sort_keys=True, default=str)
        for row in network_policy_rows
    )
    unused_network_zones = [
        str(row.get("Name") or "Unnamed Zone")
        for row in zone_rows
        if str(row.get("Name") or "").strip()
        and str(row.get("Name")).lower() not in zone_usage_blob.lower()
    ]
    empty_trusted_zones = [
        str(row.get("Name") or "Unnamed Zone")
        for row in zone_rows
        if str(row.get("Usage") or "").strip().upper() == "TRUSTED"
        and _blankish(row.get("Gateways"))
        and _blankish(row.get("Proxies"))
    ]
    validations.extend([
        _identity_validation(
            "Network Zones Defined But Not Used In Policies",
            "Low",
            unused_network_zones,
            lambda items: f"{len(items)} network zone(s) are defined but not referenced by extracted session or app sign-on policy rules.",
            "All defined network zones are referenced by extracted session or app sign-on policy rules, or no network zones are configured.",
            "No network zone data was available for validation.",
            data_available=bool(zone_rows),
        ),
        _identity_validation(
            "Trusted Network Zones Without Entries",
            "Moderate",
            empty_trusted_zones,
            lambda items: f"{len(items)} trusted network zone(s) do not list gateways or proxies.",
            "All trusted network zones contain gateway or proxy entries.",
            "No network zone data was available for trusted-zone validation.",
            data_available=bool(zone_rows),
        ),
    ])

    # Authenticators
    authenticator_rows = _section_rows(section_by_id, "authenticators")
    weak_authenticators = [
        f"{row.get('Name')}: {row.get('Status')}"
        for row in authenticator_rows
        if str(row.get("Status") or "").strip().upper() == "ACTIVE"
        and _contains_token(row.get("Key"), ["email", "phone", "security_question"])
    ]
    strong_authenticator_present = any(
        str(row.get("Status") or "").strip().upper() == "ACTIVE"
        and _contains_token(row.get("Key"), ["okta_verify", "webauthn", "security_key"])
        for row in authenticator_rows
    )
    validations.extend([
        _identity_validation(
            "Weak Authenticators Enabled",
            "High",
            weak_authenticators,
            lambda items: f"{len(items)} weaker authenticator(s) are active.",
            "No weaker authenticators were detected as active.",
            "No authenticator data was available for validation.",
            data_available=bool(authenticator_rows),
        ),
        _identity_validation(
            "Phishing-Resistant Authenticator Availability",
            "High",
            [] if strong_authenticator_present else ["No active Okta Verify / WebAuthn / security key authenticator detected"],
            lambda items: items[0],
            "At least one phishing-resistant authenticator is available.",
            "No authenticator data was available for phishing-resistant authenticator validation.",
            data_available=bool(authenticator_rows),
        ),
    ])

    # Applications
    application_rows = _section_rows(section_by_id, "applications")
    everyone_assigned_apps = [
        str(row.get("Name") or "Unnamed App")
        for row in application_rows
        if _contains_token(row.get("Groups"), ["Everyone"])
    ]
    password_based_apps = [
        str(row.get("Name") or "Unnamed App")
        for row in application_rows
        if str(row.get("Type") or "").upper() in {"AUTO_LOGIN", "BROWSER_PLUGIN", "SECURE_PASSWORD_STORE"}
    ]
    validations.extend([
        _identity_validation(
            "Applications Assigned to Everyone",
            "Moderate",
            everyone_assigned_apps,
            lambda items: f"{len(items)} application(s) are assigned to the Everyone group.",
            "No applications were found assigned to Everyone.",
            "No application data was available for assignment validation.",
            data_available=bool(application_rows),
        ),
        _identity_validation(
            "Password-Based Application Sign-On Modes Present",
            "Low",
            password_based_apps,
            lambda items: f"{len(items)} application(s) use password-based or SWA-style sign-on modes.",
            "No password-based or SWA-style application sign-on modes were detected.",
            "No application data was available for sign-on-mode validation.",
            data_available=bool(application_rows),
        ),
    ])

    # Identity providers and discovery
    idp_rows = _section_rows(section_by_id, "identity-providers")
    inactive_idps = [
        str(row.get("Name") or "Unnamed IdP")
        for row in idp_rows
        if _looks_like_inactive(row.get("Status"))
    ]
    idp_discovery_rows = _section_rows(section_by_id, "idp-discovery-policies", entry_type="Rule")
    inactive_idp_discovery_rules = [
        f"{row.get('Policy Name')}: {row.get('Rule Name')}"
        for row in idp_discovery_rows
        if _looks_like_inactive(row.get("Status"))
    ]
    validations.extend([
        _identity_validation(
            "Inactive Identity Providers",
            "Low",
            inactive_idps,
            lambda items: f"{len(items)} identity provider(s) are not active.",
            "All identity providers are active.",
            "No identity provider data was available for validation.",
            data_available=True,
        ),
        _identity_validation(
            "Inactive IdP Discovery Rules",
            "Low",
            inactive_idp_discovery_rules,
            lambda items: f"{len(items)} IdP discovery rule(s) are not active.",
            "All IdP discovery rules are active.",
            "No IdP discovery policy rule data was available for validation.",
            data_available=True,
        ),
    ])

    # Authorization server checks
    authz_server_rows = _section_rows(section_by_id, "authz-servers")
    authz_access_rule_rows = _section_rows(section_by_id, "authz-access-policies", entry_type="Rule")
    manual_rotation_servers = [
        str(row.get("Name") or "Unnamed Authorization Server")
        for row in authz_server_rows
        if str(row.get("Credentials Rotation Mode") or "").strip().upper() not in {"AUTO", "AUTOMATIC"}
    ]
    authz_any_client_rules = [
        f"{row.get('Authorization Server')} / {row.get('Policy Name')} / {row.get('Rule Name')}"
        for row in authz_access_rule_rows
        if _contains_token(row.get("Conditions"), ["Any client", "Any"])
    ]
    validations.extend([
        _identity_validation(
            "Authorization Servers Without Automatic Key Rotation",
            "Moderate",
            manual_rotation_servers,
            lambda items: f"{len(items)} authorization server(s) do not report automatic credential rotation.",
            "All authorization servers report automatic credential rotation.",
            "No authorization server settings were available for rotation validation.",
            data_available=bool(authz_server_rows),
        ),
        _identity_validation(
            "Authorization Server Access Rules With Broad Client Scope",
            "Moderate",
            authz_any_client_rules,
            lambda items: f"{len(items)} authorization server access rule(s) appear to apply to any client.",
            "No authorization server access rules were found with broad any-client scope.",
            "No authorization server access policy rule data was available for validation.",
            data_available=bool(authz_access_rule_rows),
        ),
    ])

    # Admin governance
    custom_role_rows = _section_rows(section_by_id, "custom-admin-roles")
    custom_roles_missing_description = [
        str(row.get("Label") or "Unnamed Role")
        for row in custom_role_rows
        if _blankish(row.get("Description"))
    ]
    resource_set_rows = _section_rows(section_by_id, "resource-sets")
    resource_sets = {}
    for row in resource_set_rows:
        entry_type = str(row.get("Entry Type") or "")
        label = row.get("Label") or row.get("Resource Set")
        if not label:
            continue
        bucket = resource_sets.setdefault(str(label), {"resource": 0, "binding": 0})
        if entry_type == "Resource":
            bucket["resource"] += 1
        elif entry_type == "Binding":
            bucket["binding"] += 1
    resource_sets_without_resources = [name for name, counts in sorted(resource_sets.items()) if counts["resource"] == 0]
    resource_sets_without_bindings = [name for name, counts in sorted(resource_sets.items()) if counts["binding"] == 0]
    admin_app_rows = _section_rows(section_by_id, "admin-assignments-apps")
    validations.extend([
        _identity_validation(
            "Custom Admin Roles Missing Description",
            "Low",
            custom_roles_missing_description,
            lambda items: f"{len(items)} custom admin role(s) do not have a description.",
            "All custom admin roles have descriptions.",
            "No custom admin role data was available for validation.",
            data_available=bool(custom_role_rows),
        ),
        _identity_validation(
            "Resource Sets Without Resources",
            "Moderate",
            resource_sets_without_resources,
            lambda items: f"{len(items)} resource set(s) do not contain any resources.",
            "All resource sets contain at least one resource.",
            "No resource set data was available for validation.",
            data_available=bool(resource_set_rows),
        ),
        _identity_validation(
            "Resource Sets Without Bindings",
            "Moderate",
            resource_sets_without_bindings,
            lambda items: f"{len(items)} resource set(s) do not contain any bindings.",
            "All resource sets contain at least one binding.",
            "No resource set data was available for validation.",
            data_available=bool(resource_set_rows),
        ),
        _identity_validation(
            "Admin Public Client Applications Present",
            "Moderate",
            [str(row.get("Display Name") or row.get("App Name") or "Unnamed Admin App") for row in admin_app_rows],
            lambda items: f"{len(items)} admin public client application(s) are configured and should be reviewed.",
            "No admin public client applications were detected.",
            "No admin app assignment data was available for validation.",
            data_available=bool(admin_app_rows) or bool(_section_rows(section_by_id, "admin-assignments-apps")),
        ),
    ])

    api_token_rows = _section_rows(section_by_id, "api-tokens")
    tokens_without_network_restrictions = []
    for row in api_token_rows:
        network = row.get("Network")
        if isinstance(network, dict):
            connection = str(network.get("connection") or "").strip().upper()
            include = network.get("include") or []
            exclude = network.get("exclude") or []
            if connection in {"", "ANYWHERE", "ANY_NETWORK"} and not include and not exclude:
                tokens_without_network_restrictions.append(str(row.get("Name") or row.get("Token ID") or "Unnamed Token"))
        elif _blankish(network):
            tokens_without_network_restrictions.append(str(row.get("Name") or row.get("Token ID") or "Unnamed Token"))
    validations.append(_identity_validation(
        "API Tokens Without Network Restrictions",
        "High",
        tokens_without_network_restrictions,
        lambda items: f"{len(items)} API token(s) do not show network restrictions.",
        "All extracted API tokens show some form of network restriction.",
        "No API token data was available for network restriction validation.",
        data_available=bool(api_token_rows),
    ))

    # Brand and UX
    brand_setting_rows = _section_rows(section_by_id, "brand-settings")
    brands_missing_privacy = [
        str(row.get("Brand Name") or "Unnamed Brand")
        for row in brand_setting_rows
        if _blankish(row.get("Custom Privacy Policy URL"))
    ]
    validations.extend([
        _identity_validation(
            "Brands Missing Custom Privacy Policy URL",
            "Low",
            brands_missing_privacy,
            lambda items: f"{len(items)} brand(s) do not define a custom privacy policy URL.",
            "All brands define a custom privacy policy URL.",
            "No brand settings were available for validation.",
            data_available=bool(brand_setting_rows),
        ),
    ])

    # Trusted origins and hooks
    trusted_origin_rows = _section_rows(section_by_id, "trusted-origins")
    insecure_origins = [
        str(row.get("Origin") or row.get("Name") or "Unnamed Origin")
        for row in trusted_origin_rows
        if str(row.get("Origin") or "").strip().lower().startswith("http://")
    ]
    event_hook_rows = _section_rows(section_by_id, "event-hooks")
    unverified_event_hooks = [
        str(row.get("Name") or "Unnamed Event Hook")
        for row in event_hook_rows
        if str(row.get("Verification Status") or "").strip().upper() != "VERIFIED"
    ]
    weak_event_hook_auth = [
        str(row.get("Name") or "Unnamed Event Hook")
        for row in event_hook_rows
        if _blankish(row.get("Auth Scheme"))
    ]
    inactive_event_hooks = [
        str(row.get("Name") or "Unnamed Event Hook")
        for row in event_hook_rows
        if _looks_like_inactive(row.get("Status"))
    ]
    inline_hook_rows = _section_rows(section_by_id, "inline-hooks")
    inactive_inline_hooks = [
        str(row.get("Name") or "Unnamed Inline Hook")
        for row in inline_hook_rows
        if _looks_like_inactive(row.get("Status"))
    ]
    weak_inline_hook_auth = [
        str(row.get("Name") or "Unnamed Inline Hook")
        for row in inline_hook_rows
        if _blankish(row.get("Auth Scheme"))
    ]
    validations.extend([
        _identity_validation(
            "Insecure Trusted Origins",
            "High",
            insecure_origins,
            lambda items: f"{len(items)} trusted origin(s) use HTTP instead of HTTPS.",
            "No HTTP trusted origins were detected.",
            "No trusted origin data was available for protocol validation.",
            data_available=bool(trusted_origin_rows),
        ),
        _identity_validation(
            "Unverified Event Hooks",
            "High",
            unverified_event_hooks,
            lambda items: f"{len(items)} event hook(s) are not verified.",
            "All event hooks are verified, or no event hooks are configured.",
            "No event hook data was available for verification-status validation.",
            data_available=True,
        ),
        _identity_validation(
            "Event Hooks Without Authentication Scheme",
            "High",
            weak_event_hook_auth,
            lambda items: f"{len(items)} event hook(s) do not report an authentication scheme.",
            "All event hooks report an authentication scheme, or no event hooks are configured.",
            "No event hook data was available for auth-scheme validation.",
            data_available=True,
        ),
        _identity_validation(
            "Inactive Event Hooks",
            "Low",
            inactive_event_hooks,
            lambda items: f"{len(items)} event hook(s) are not active.",
            "All event hooks are active, or no event hooks are configured.",
            "No event hook data was available for status validation.",
            data_available=True,
        ),
        _identity_validation(
            "Inline Hooks Without Authentication Scheme",
            "High",
            weak_inline_hook_auth,
            lambda items: f"{len(items)} inline hook(s) do not report an authentication scheme.",
            "All inline hooks report an authentication scheme.",
            "No inline hook data was available for auth-scheme validation.",
            data_available=bool(inline_hook_rows),
        ),
        _identity_validation(
            "Inactive Inline Hooks",
            "Low",
            inactive_inline_hooks,
            lambda items: f"{len(items)} inline hook(s) are not active.",
            "All inline hooks are active, or no inline hooks are configured.",
            "No inline hook data was available for status validation.",
            data_available=True,
        ),
    ])

    # Attack protection
    attack_rows = _section_rows(section_by_id, "attack-protection")
    weak_attack_protection_rows = [
        f"{row.get('Component')} / {row.get('Field')}: {row.get('Value')}"
        for row in attack_rows
        if _contains_token(row.get("Field"), ["enabled", "mode", "status", "action"])
        and _contains_token(row.get("Value"), ["disabled", "false", "none", "log"])
    ]
    validations.append(_identity_validation(
        "Attack Protection Controls Not Enforcing",
        "High",
        weak_attack_protection_rows,
        lambda items: f"{len(items)} attack-protection setting(s) appear disabled or monitor-only.",
        "No extracted attack-protection settings appeared disabled or monitor-only.",
        "No attack-protection data was available for validation.",
        data_available=bool(attack_rows),
    ))

    # Realms
    realm_rows = _section_rows(section_by_id, "realms")
    inactive_realms = [
        str(row.get("Name") or "Unnamed Realm")
        for row in realm_rows
        if _looks_like_inactive(row.get("Status"))
    ]
    realm_assignment_rows = _section_rows(section_by_id, "realm-assignments")
    default_realm_assignments = [
        str(row.get("Name") or "Unnamed Assignment")
        for row in realm_assignment_rows
        if str(row.get("Is Default") or "").strip().lower() == "true"
    ]
    validations.extend([
        _identity_validation(
            "Inactive Realms",
            "Low",
            inactive_realms,
            lambda items: f"{len(items)} realm(s) are not active.",
            "All realms are active.",
            "No realm data was available for validation.",
            data_available=True,
        ),
        _identity_validation(
            "Realm Default Assignment Coverage",
            "High",
            [] if len(default_realm_assignments) == 1 else [f"Default assignments found: {len(default_realm_assignments)}"],
            lambda items: "Realm assignments should contain exactly one default assignment.",
            "Realm assignments contain exactly one default assignment.",
            "No realm assignment data was available for default-assignment validation.",
            data_available=bool(realm_assignment_rows),
        ),
    ])

    # Profile schema and mappings
    schema_rows = _section_rows(section_by_id, "profile-schema-user")
    sensitive_writable_attrs = [
        f"{row.get('Attribute')}: {row.get('Mutability')}"
        for row in schema_rows
        if _contains_token(row.get("Attribute"), ["password", "secret", "token", "ssn", "salary"])
        and _contains_token(row.get("Mutability"), ["read_write", "write"])
    ]
    profile_mapping_rows = _section_rows(section_by_id, "profile-mappings")
    empty_profile_mappings = [
        f"{row.get('Source Name')} -> {row.get('Target Name')}"
        for row in profile_mapping_rows
        if _blankish(row.get("Properties")) or _text(row.get("Properties")) == "{}"
    ]
    validations.extend([
        _identity_validation(
            "Sensitive Writable Profile Attributes",
            "High",
            sensitive_writable_attrs,
            lambda items: f"{len(items)} sensitive profile attribute(s) appear writable.",
            "No sensitive writable profile attributes were detected by name heuristic.",
            "No user profile schema data was available for validation.",
            data_available=bool(schema_rows),
        ),
        _identity_validation(
            "Empty Profile Mappings",
            "Moderate",
            empty_profile_mappings,
            lambda items: f"{len(items)} profile mapping(s) do not contain extracted property mappings.",
            "All extracted profile mappings contain property mappings.",
            "No profile mapping data was available for validation.",
            data_available=bool(profile_mapping_rows),
        ),
    ])

    # Group push mappings
    group_push_rows = _section_rows(section_by_id, "group-push-mappings")
    inactive_group_push = [
        f"{row.get('App Name')}: {row.get('Source Group')} -> {row.get('Target Group')}"
        for row in group_push_rows
        if _looks_like_inactive(row.get("Status"))
    ]
    stale_group_push = [
        f"{row.get('App Name')}: {row.get('Source Group')} -> {row.get('Target Group')}"
        for row in group_push_rows
        if (_days_since(row.get("Last Updated")) or 0) >= 90
    ]
    validations.extend([
        _identity_validation(
            "Inactive Group Push Mappings",
            "Low",
            inactive_group_push,
            lambda items: f"{len(items)} group push mapping(s) are not active.",
            "All group push mappings are active, or no group push mappings are configured.",
            "No group push mapping data was available for validation.",
            data_available=True,
        ),
        _identity_validation(
            "Stale Group Push Mappings",
            "Low",
            stale_group_push,
            lambda items: f"{len(items)} group push mapping(s) have not been updated in at least 90 days.",
            "No stale group push mappings were detected using the 90-day age heuristic.",
            "No group push mapping data was available for age validation.",
            data_available=bool(group_push_rows),
        ),
    ])

    # Risk-aware policies
    entity_risk_rule_rows = _section_rows(section_by_id, "entity-risk-policies", entry_type="Rule")
    post_auth_rule_rows = _section_rows(section_by_id, "post-auth-session-policies", entry_type="Rule")
    active_entity_risk_rules = [
        row for row in entity_risk_rule_rows
        if not _looks_like_inactive(row.get("Status"))
    ]
    active_post_auth_rules = [
        row for row in post_auth_rule_rows
        if not _looks_like_inactive(row.get("Status"))
    ]
    validations.extend([
        _identity_validation(
            "Entity Risk Policy Rule Coverage",
            "High",
            [] if active_entity_risk_rules else ["No active entity risk policy rules detected"],
            lambda items: items[0],
            "Active entity risk policy rules were detected.",
            "No entity risk policy data was available for validation.",
            data_available=bool(_section_rows(section_by_id, "entity-risk-policies")),
        ),
        _identity_validation(
            "Identity Threat Protection Policy Rule Coverage",
            "High",
            [] if active_post_auth_rules else ["No active identity threat protection policy rules detected"],
            lambda items: items[0],
            "Active identity threat protection policy rules were detected.",
            "No identity threat protection policy data was available for validation.",
            data_available=bool(_section_rows(section_by_id, "post-auth-session-policies")),
        ),
    ])

    # HealthInsight-specific policy heuristics
    auth_policy_rule_rows = _section_rows(section_by_id, "authentication-policies", entry_type="Rule")
    session_policy_rule_rows = _section_rows(section_by_id, "global-session-policies", entry_type="Rule")
    risk_based_mfa_gaps = []
    new_device_mfa_gaps = []
    for row in auth_policy_rule_rows + session_policy_rule_rows:
        conditions = {
            "people": row.get("Conditions People"),
            "network": row.get("Conditions Network"),
            "authContext": row.get("Conditions AuthContext"),
            "risk": row.get("Conditions Risk"),
            "riskScore": row.get("Conditions RiskScore"),
            "identityProvider": row.get("Conditions IdentityProvider"),
        }
        actions = _as_dict(row.get("Actions"))
        if _conditions_reference_high_risk(conditions) and not (
            _actions_reference_mfa(actions) and _actions_reference_every_sign_in(actions)
        ):
            risk_based_mfa_gaps.append(f"{row.get('Policy Name')} / {row.get('Rule Name')}")
        if _conditions_reference_new_device(conditions) and not (
            _actions_reference_mfa(actions) and _actions_reference_every_sign_in(actions)
        ):
            new_device_mfa_gaps.append(f"{row.get('Policy Name')} / {row.get('Rule Name')}")

    admin_console_apps = [
        row for row in application_rows
        if str(row.get("Name") or "").strip().lower() == "okta admin console"
    ]
    admin_console_mfa_gap = []
    if admin_console_apps:
        for app_row in admin_console_apps:
            policy_name = str(app_row.get("Access Policy Name") or "").strip()
            if not policy_name:
                admin_console_mfa_gap.append(
                    f"{app_row.get('Name')}: no app sign-on policy is assigned"
                )
                continue
            matching_rules = [
                row for row in auth_policy_rule_rows
                if str(row.get("Policy Name") or "").strip() == policy_name
            ]
            if not matching_rules:
                admin_console_mfa_gap.append(
                    f"{app_row.get('Name')}: no app sign-on policy rules found for assigned policy '{policy_name}'"
                )
                continue
            if not any(
                _actions_reference_mfa(_as_dict(rule.get("Actions")))
                and _actions_reference_every_sign_in(_as_dict(rule.get("Actions")))
                for rule in matching_rules
                if not _looks_like_inactive(rule.get("Status"))
            ):
                admin_console_mfa_gap.append(
                    f"{app_row.get('Name')}: assigned policy '{policy_name}' has no active rule requiring MFA at every sign-in"
                )
    else:
        admin_console_mfa_gap.append("Okta Admin Console application was not found in the extracted active applications")

    validations.extend([
        _identity_validation(
            "High-Risk Request MFA Every Sign-In",
            "High",
            risk_based_mfa_gaps,
            lambda items: f"{len(items)} policy rule(s) reference risk conditions without clear MFA-at-every-sign-in actions.",
            "All extracted high-risk policy rules appear to require MFA at every sign-in.",
            "No policy rule data was available to evaluate high-risk request MFA posture.",
            data_available=bool(auth_policy_rule_rows or session_policy_rule_rows),
        ),
        _identity_validation(
            "New Device MFA Every Sign-In",
            "Moderate",
            new_device_mfa_gaps,
            lambda items: f"{len(items)} policy rule(s) reference device/new-device conditions without clear MFA-at-every-sign-in actions.",
            "All extracted device/new-device policy rules appear to require MFA at every sign-in.",
            "No policy rule data was available to evaluate new-device MFA posture.",
            data_available=bool(auth_policy_rule_rows or session_policy_rule_rows),
        ),
        _identity_validation(
            "Admin Console MFA Every Sign-In",
            "High",
            admin_console_mfa_gap,
            lambda items: f"{len(items)} admin-console application policy path(s) do not show clear MFA-at-every-sign-in enforcement.",
            "Admin Console application policies appear to require MFA at every sign-in.",
            "Admin Console application or authentication policy data was not available for validation.",
            data_available=bool(application_rows and auth_policy_rule_rows),
        ),
    ])

    # Identity-risk detections aligned to Okta ISPM where data is available from native Okta APIs
    users = extra_context.get("all_users") or []
    api_tokens = extra_context.get("api_tokens") or []
    custom_admin_roles = extra_context.get("custom_admin_roles") or []
    resource_set_bindings = extra_context.get("resource_set_bindings") or []
    user_inventory_available = "all_users" in extra_context
    token_inventory_available = "api_tokens" in extra_context
    admin_role_inventory_available = (
        "custom_admin_roles" in extra_context and "resource_set_bindings" in extra_context
    )
    token_map = _token_inventory(api_tokens)
    mfa_enforcement = _tenant_mfa_enforcement_gaps(section_by_id)

    no_mfa_accounts = []
    no_mfa_admin_accounts = []
    no_mfa_global_admin_accounts = []
    pending_mfa_accounts = []
    pending_mfa_admin_accounts = []
    pending_mfa_global_admin_accounts = []
    no_mfa_enforced_admin_accounts = []
    no_mfa_enforced_global_admin_accounts = []
    old_password_accounts = []
    old_password_service_accounts = []
    old_password_admin_accounts = []
    old_password_admin_service_accounts = []
    old_password_global_admin_accounts = []
    old_password_global_admin_service_accounts = []
    unused_accounts = []
    unused_service_accounts = []
    unused_admin_accounts = []
    unused_admin_service_accounts = []
    unused_global_admin_accounts = []
    unused_global_admin_service_accounts = []
    sso_bypass_accounts = []
    sso_bypass_admin_accounts = []
    sso_bypass_no_mfa_accounts = []
    sso_bypass_no_mfa_admin_accounts = []
    old_no_mfa_unused_accounts = []
    old_no_mfa_unused_admin_accounts = []
    old_no_mfa_unused_service_accounts = []
    old_no_mfa_unused_admin_service_accounts = []
    partially_offboarded_users = []
    super_admin_with_api_token = []
    service_account_console_access_accounts = []
    service_account_console_access_admin_accounts = []
    unrotated_tokens_account = []
    unrotated_tokens_admin = []
    unused_tokens_account = []
    unused_tokens_admin = []
    unrotated_unused_tokens_account = []
    unrotated_unused_tokens_admin = []
    super_admins = []

    for user in users:
        user_id = user.get("id")
        label = _user_identifier(user)
        role_types = _user_role_types(user)
        is_admin = bool(role_types)
        is_super_admin = "SUPER_ADMIN" in role_types
        active_tokens = token_map.get(user_id, [])
        token_count = len(active_tokens)
        is_service_account = _is_service_account(user, active_token_count=token_count)
        has_mfa = _has_mfa(user)
        pending_mfa = _has_pending_mfa(user)
        direct_access = _is_direct_access_user(user)
        old_password = _has_old_password(user, threshold_days=90)
        unused_user = _is_unused_user(user, threshold_days=91)
        no_mfa = not has_mfa
        status = str(user.get("status") or "").strip().upper()
        org_mfa_gap = mfa_enforcement["has_data"] and (
            mfa_enforcement["optional_enrollment"] or not mfa_enforcement["auth_requires_mfa"]
        )

        if is_super_admin:
            super_admins.append(label)

        if no_mfa:
            no_mfa_accounts.append(label)
            if is_admin:
                no_mfa_admin_accounts.append(label)
            if is_super_admin:
                no_mfa_global_admin_accounts.append(label)

        if pending_mfa:
            pending_mfa_accounts.append(f"{label}: pending factors {', '.join(sorted(_pending_factor_types(user)))}")
            if is_admin:
                pending_mfa_admin_accounts.append(f"{label}: pending factors {', '.join(sorted(_pending_factor_types(user)))}")
            if is_super_admin:
                pending_mfa_global_admin_accounts.append(f"{label}: pending factors {', '.join(sorted(_pending_factor_types(user)))}")

        if old_password:
            old_password_accounts.append(f"{label}: password age {_days_since(user.get('passwordChanged'))}d")
            if is_service_account:
                old_password_service_accounts.append(f"{label}: password age {_days_since(user.get('passwordChanged'))}d")
            if is_admin:
                old_password_admin_accounts.append(f"{label}: password age {_days_since(user.get('passwordChanged'))}d")
            if is_admin and is_service_account:
                old_password_admin_service_accounts.append(f"{label}: password age {_days_since(user.get('passwordChanged'))}d")
            if is_super_admin:
                old_password_global_admin_accounts.append(f"{label}: password age {_days_since(user.get('passwordChanged'))}d")
            if is_super_admin and is_service_account:
                old_password_global_admin_service_accounts.append(f"{label}: password age {_days_since(user.get('passwordChanged'))}d")

        if unused_user:
            days_text = _days_since(user.get("lastLogin"))
            suffix = f"last login {days_text}d ago" if days_text is not None else "no recent interactive login"
            unused_accounts.append(f"{label}: {suffix}")
            if is_service_account:
                unused_service_accounts.append(f"{label}: {suffix}")
            if is_admin:
                unused_admin_accounts.append(f"{label}: {suffix}")
            if is_admin and is_service_account:
                unused_admin_service_accounts.append(f"{label}: {suffix}")
            if is_super_admin:
                unused_global_admin_accounts.append(f"{label}: {suffix}")
            if is_super_admin and is_service_account:
                unused_global_admin_service_accounts.append(f"{label}: {suffix}")

        if direct_access:
            sso_bypass_accounts.append(f"{label}: provider {_user_provider_type(user) or 'unknown'}")
            if is_admin:
                sso_bypass_admin_accounts.append(f"{label}: provider {_user_provider_type(user) or 'unknown'}")

        if direct_access and no_mfa:
            sso_bypass_no_mfa_accounts.append(f"{label}: provider {_user_provider_type(user) or 'unknown'}")
            if is_admin:
                sso_bypass_no_mfa_admin_accounts.append(f"{label}: provider {_user_provider_type(user) or 'unknown'}")

        if old_password and no_mfa and unused_user:
            old_no_mfa_unused_accounts.append(label)
            if is_admin:
                old_no_mfa_unused_admin_accounts.append(label)
            if is_service_account:
                old_no_mfa_unused_service_accounts.append(label)
            if is_admin and is_service_account:
                old_no_mfa_unused_admin_service_accounts.append(label)

        if org_mfa_gap and is_admin:
            no_mfa_enforced_admin_accounts.append(label)
        if org_mfa_gap and is_super_admin:
            no_mfa_enforced_global_admin_accounts.append(label)

        if status in {"SUSPENDED", "DEPROVISIONED"} and (is_admin or active_tokens or direct_access):
            indicators = []
            if is_admin:
                indicators.append("admin role assignment")
            if active_tokens:
                indicators.append(f"{len(active_tokens)} active API token(s)")
            if direct_access:
                indicators.append(f"provider {_user_provider_type(user) or 'unknown'}")
            partially_offboarded_users.append(f"{label}: status {status}; remaining access via {', '.join(indicators)}")

        if is_service_account and direct_access:
            service_account_console_access_accounts.append(f"{label}: provider {_user_provider_type(user) or 'unknown'}")
            if is_admin:
                service_account_console_access_admin_accounts.append(f"{label}: provider {_user_provider_type(user) or 'unknown'}")

        if is_super_admin and active_tokens:
            for token in active_tokens:
                age_days = _token_rotation_age_days(token)
                suffix = f"age {age_days}d" if age_days is not None else "active token"
                super_admin_with_api_token.append(_format_token_item(user, token, suffix))

        unrotated_tokens, unused_tokens = _token_findings_for_user(active_tokens, owner_unused=unused_user, rotation_days=90)
        for token, age_days in unrotated_tokens:
            item = _format_token_item(user, token, f"rotation age {age_days}d")
            if is_admin:
                unrotated_tokens_admin.append(item)
            else:
                unrotated_tokens_account.append(item)
        for token in unused_tokens:
            item = _format_token_item(user, token, "owner inactive >=91d")
            if is_admin:
                unused_tokens_admin.append(item)
            else:
                unused_tokens_account.append(item)
        for token, age_days in unrotated_tokens:
            if token in unused_tokens:
                item = _format_token_item(user, token, f"rotation age {age_days}d; owner inactive >=91d")
                if is_admin:
                    unrotated_unused_tokens_admin.append(item)
                else:
                    unrotated_unused_tokens_account.append(item)

    role_usage = _custom_role_usage_map(custom_admin_roles, resource_set_bindings)
    unused_custom_roles = [
        str(
            details["role"].get("label")
            or details["role"].get("name")
            or details["role"].get("id")
            or "Unnamed Role"
        )
        for details in role_usage.values()
        if not details.get("used")
    ]

    validations.extend([
        _identity_validation(
            "Old Password, No MFA, Unused Accounts",
            "High",
            old_no_mfa_unused_accounts,
            lambda items: f"{len(items)} account(s) have passwords older than 90 days, no active MFA, and no recent interactive use for at least 91 days.",
            "No accounts matched the old-password, no-MFA, unused-account combination heuristic.",
            "User inventory was not available to assess old-password, no-MFA, unused-account combinations.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Old Password, No MFA, Unused Admin Accounts",
            "High",
            old_no_mfa_unused_admin_accounts,
            lambda items: f"{len(items)} admin account(s) have passwords older than 90 days, no active MFA, and no recent interactive use for at least 91 days.",
            "No admin accounts matched the old-password, no-MFA, unused-account combination heuristic.",
            "User inventory was not available to assess old-password, no-MFA, unused admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Old Password, No MFA, Unused Service Accounts",
            "High",
            old_no_mfa_unused_service_accounts,
            lambda items: f"{len(items)} potential service account(s) have old passwords, no active MFA, and no recent interactive use.",
            "No potential service accounts matched the old-password, no-MFA, unused heuristic.",
            "User inventory was not available to assess old-password, no-MFA, unused service accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Old Password, No MFA, Unused Admin Service Accounts",
            "High",
            old_no_mfa_unused_admin_service_accounts,
            lambda items: f"{len(items)} potential admin service account(s) have old passwords, no active MFA, and no recent interactive use.",
            "No potential admin service accounts matched the old-password, no-MFA, unused heuristic.",
            "User inventory was not available to assess old-password, no-MFA, unused admin service accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "SSO Bypass + No MFA - Admin",
            "High",
            sso_bypass_no_mfa_admin_accounts,
            lambda items: f"{len(items)} admin account(s) appear to have direct access that can bypass federated SSO and have no active MFA.",
            "No admin accounts matched the direct-access plus no-active-MFA heuristic.",
            "User inventory was not available to assess SSO bypass plus no MFA for admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "No MFA - Global Admin Account",
            "High",
            no_mfa_global_admin_accounts,
            lambda items: f"{len(items)} super admin account(s) do not have an active enrolled factor.",
            "All detected super admin accounts have at least one active enrolled factor.",
            "User inventory was not available to assess MFA enrollment for super admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Pending MFA - Global Admin Account",
            "High",
            pending_mfa_global_admin_accounts,
            lambda items: f"{len(items)} super admin account(s) appear to have pending factor enrollment without an active enrolled factor.",
            "No super admin accounts were found with only pending factor enrollment.",
            "User inventory was not available to assess pending MFA for super admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Old Password - Global Admin Account",
            "High",
            old_password_global_admin_accounts,
            lambda items: f"{len(items)} super admin account(s) have passwords older than 90 days.",
            "No super admin accounts were found with passwords older than 90 days.",
            "User inventory was not available to assess password age for super admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Old Password - Global Admin Service Account",
            "High",
            old_password_global_admin_service_accounts,
            lambda items: f"{len(items)} potential super admin service account(s) have passwords older than 90 days.",
            "No potential super admin service accounts were found with passwords older than 90 days.",
            "User inventory was not available to assess password age for super admin service accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Unused Global Admin Account",
            "High",
            unused_global_admin_accounts,
            lambda items: f"{len(items)} super admin account(s) have not logged in interactively for at least 91 days.",
            "No unused super admin accounts were detected using the 91-day inactivity heuristic.",
            "User inventory was not available to assess unused super admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Unused Global Admin Service Account",
            "High",
            unused_global_admin_service_accounts,
            lambda items: f"{len(items)} potential super admin service account(s) have not logged in interactively for at least 91 days.",
            "No unused potential super admin service accounts were detected using the 91-day inactivity heuristic.",
            "User inventory was not available to assess unused super admin service accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Partially Off-boarded User",
            "High",
            partially_offboarded_users,
            lambda items: f"{len(items)} suspended or deprovisioned user(s) still appear to retain direct access, admin roles, or active API tokens.",
            "No partially off-boarded users were detected with remaining privileged or direct access indicators.",
            "User inventory was not available to assess partially off-boarded users.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "No MFA - Admin Account",
            "High",
            no_mfa_admin_accounts,
            lambda items: f"{len(items)} admin account(s) do not have an active enrolled factor.",
            "All detected admin accounts have at least one active enrolled factor.",
            "User inventory was not available to assess MFA enrollment for admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Pending MFA - Admin Account",
            "High",
            pending_mfa_admin_accounts,
            lambda items: f"{len(items)} admin account(s) appear to have pending factor enrollment without an active enrolled factor.",
            "No admin accounts were found with only pending factor enrollment.",
            "User inventory was not available to assess pending MFA for admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "No MFA Enforced - Global Admin Account",
            "High",
            no_mfa_enforced_global_admin_accounts if mfa_enforcement["has_data"] else [],
            lambda items: f"{len(items)} super admin account(s) may not be covered by enforced MFA requirements based on tenant-wide enrollment/authentication policy heuristics.",
            "Tenant-wide policy heuristics suggest MFA enforcement is present for super admin accounts.",
            "Policy data was insufficient to infer tenant-wide MFA enforcement posture for super admin accounts.",
            data_available=mfa_enforcement["has_data"],
        ),
        _identity_validation(
            "Excessive Number of Super Admins",
            "High",
            super_admins if len(super_admins) > 3 else [],
            lambda items: f"{len(items)} super admin accounts were detected. The heuristic threshold is more than 3.",
            f"{len(super_admins)} super admin account(s) were detected, which is within the heuristic threshold of 3.",
            "User inventory was not available to assess the number of super admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "SSO Bypass - Admin",
            "High",
            sso_bypass_admin_accounts,
            lambda items: f"{len(items)} admin account(s) appear to support direct access outside federated SSO based on their credential provider.",
            "No admin accounts matched the direct-access credential provider heuristic.",
            "User inventory was not available to assess SSO bypass for admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "SSO Bypass + No MFA - Account",
            "High",
            sso_bypass_no_mfa_accounts,
            lambda items: f"{len(items)} account(s) appear to support direct access outside federated SSO and have no active MFA.",
            "No accounts matched the direct-access plus no-active-MFA heuristic.",
            "User inventory was not available to assess SSO bypass plus no MFA for accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Old Password - Admin Account",
            "High",
            old_password_admin_accounts,
            lambda items: f"{len(items)} admin account(s) have passwords older than 90 days.",
            "No admin accounts were found with passwords older than 90 days.",
            "User inventory was not available to assess password age for admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Old Password - Admin Service Account",
            "High",
            old_password_admin_service_accounts,
            lambda items: f"{len(items)} potential admin service account(s) have passwords older than 90 days.",
            "No potential admin service accounts were found with passwords older than 90 days.",
            "User inventory was not available to assess password age for admin service accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Super Admin with API Token",
            "High",
            super_admin_with_api_token,
            lambda items: f"{len(items)} active API token(s) are owned by super admin accounts.",
            "No active API tokens owned by super admin accounts were detected.",
            "API token or user inventory was not available to assess super admin token ownership.",
            data_available=user_inventory_available and token_inventory_available,
        ),
        _identity_validation(
            "Unrotated and Unused Keys and Tokens - Admin",
            "High",
            unrotated_unused_tokens_admin,
            lambda items: f"{len(items)} active API token(s) owned by admin accounts are older than 90 days and the owning account has been inactive for at least 91 days.",
            "No admin-owned active API tokens matched the unrotated-and-owner-inactive heuristic.",
            "API token or user inventory was not available to assess unrotated and unused admin tokens.",
            data_available=user_inventory_available and token_inventory_available,
        ),
        _identity_validation(
            "Unrotated Keys and Tokens - Admin",
            "High",
            unrotated_tokens_admin,
            lambda items: f"{len(items)} active API token(s) owned by admin accounts are older than 90 days.",
            "No admin-owned active API tokens older than 90 days were detected.",
            "API token or user inventory was not available to assess token rotation age for admin accounts.",
            data_available=user_inventory_available and token_inventory_available,
        ),
        _identity_validation(
            "Unused Admin Account",
            "High",
            unused_admin_accounts,
            lambda items: f"{len(items)} admin account(s) have not logged in interactively for at least 91 days.",
            "No unused admin accounts were detected using the 91-day inactivity heuristic.",
            "User inventory was not available to assess unused admin accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Unused Admin Service Account",
            "High",
            unused_admin_service_accounts,
            lambda items: f"{len(items)} potential admin service account(s) have not logged in interactively for at least 91 days.",
            "No unused potential admin service accounts were detected using the 91-day inactivity heuristic.",
            "User inventory was not available to assess unused admin service accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Unused Keys and Tokens - Admins",
            "High",
            unused_tokens_admin,
            lambda items: f"{len(items)} active API token(s) are owned by admin accounts that have been inactive for at least 91 days.",
            "No admin-owned active API tokens matched the owner-inactive heuristic.",
            "API token or user inventory was not available to assess unused admin tokens.",
            data_available=user_inventory_available and token_inventory_available,
        ),
        _identity_validation(
            "Service Account with Console Access - Admin",
            "High",
            service_account_console_access_admin_accounts,
            lambda items: f"{len(items)} potential admin service account(s) appear to have direct console access based on their credential provider.",
            "No potential admin service accounts matched the direct-console-access heuristic.",
            "User inventory was not available to assess direct console access for admin service accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Unused Administrative Roles",
            "High",
            unused_custom_roles,
            lambda items: f"{len(items)} custom admin role(s) were found without a matching resource-set binding.",
            "All discovered custom admin roles were referenced by at least one resource-set binding.",
            "Custom admin role inventory was not available to assess unused administrative roles.",
            data_available=admin_role_inventory_available,
        ),
        _identity_validation(
            "No MFA Enforced - Admin Account",
            "Moderate",
            no_mfa_enforced_admin_accounts if mfa_enforcement["has_data"] else [],
            lambda items: f"{len(items)} admin account(s) may not be covered by enforced MFA requirements based on tenant-wide enrollment/authentication policy heuristics.",
            "Tenant-wide policy heuristics suggest MFA enforcement is present for admin accounts.",
            "Policy data was insufficient to infer tenant-wide MFA enforcement posture for admin accounts.",
            data_available=mfa_enforcement["has_data"],
        ),
        _identity_validation(
            "No MFA - Account",
            "Moderate",
            no_mfa_accounts,
            lambda items: f"{len(items)} account(s) do not have an active enrolled factor.",
            "All assessed accounts have at least one active enrolled factor.",
            "User inventory was not available to assess MFA enrollment for accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Pending MFA",
            "Moderate",
            pending_mfa_accounts,
            lambda items: f"{len(items)} account(s) appear to have pending factor enrollment without an active enrolled factor.",
            "No accounts were found with only pending factor enrollment.",
            "User inventory was not available to assess pending MFA for accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "SSO Bypass - Account",
            "Moderate",
            sso_bypass_accounts,
            lambda items: f"{len(items)} account(s) appear to support direct access outside federated SSO based on their credential provider.",
            "No accounts matched the direct-access credential provider heuristic.",
            "User inventory was not available to assess SSO bypass for accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Old Password - Account",
            "Moderate",
            old_password_accounts,
            lambda items: f"{len(items)} account(s) have passwords older than 90 days.",
            "No accounts were found with passwords older than 90 days.",
            "User inventory was not available to assess password age for accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Old Password - Service Account",
            "Moderate",
            old_password_service_accounts,
            lambda items: f"{len(items)} potential service account(s) have passwords older than 90 days.",
            "No potential service accounts were found with passwords older than 90 days.",
            "User inventory was not available to assess password age for potential service accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Unrotated and Unused Keys and Tokens - Account",
            "Moderate",
            unrotated_unused_tokens_account,
            lambda items: f"{len(items)} active API token(s) owned by non-admin accounts are older than 90 days and the owning account has been inactive for at least 91 days.",
            "No non-admin active API tokens matched the unrotated-and-owner-inactive heuristic.",
            "API token or user inventory was not available to assess unrotated and unused tokens for accounts.",
            data_available=user_inventory_available and token_inventory_available,
        ),
        _identity_validation(
            "Unrotated Keys and Tokens - Account",
            "Moderate",
            unrotated_tokens_account,
            lambda items: f"{len(items)} active API token(s) owned by non-admin accounts are older than 90 days.",
            "No non-admin active API tokens older than 90 days were detected.",
            "API token or user inventory was not available to assess token rotation age for accounts.",
            data_available=user_inventory_available and token_inventory_available,
        ),
        _identity_validation(
            "Unused Account",
            "Low",
            unused_accounts,
            lambda items: f"{len(items)} account(s) have not logged in interactively for at least 91 days.",
            "No unused accounts were detected using the 91-day inactivity heuristic.",
            "User inventory was not available to assess unused accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Unused Service Account",
            "Low",
            unused_service_accounts,
            lambda items: f"{len(items)} potential service account(s) have not logged in interactively for at least 91 days.",
            "No unused potential service accounts were detected using the 91-day inactivity heuristic.",
            "User inventory was not available to assess unused service accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Service Account with Console Access - Account",
            "Low",
            service_account_console_access_accounts,
            lambda items: f"{len(items)} potential service account(s) appear to have direct console access based on their credential provider.",
            "No potential service accounts matched the direct-console-access heuristic.",
            "User inventory was not available to assess direct console access for service accounts.",
            data_available=user_inventory_available,
        ),
        _identity_validation(
            "Unused Keys and Tokens - Account",
            "Low",
            unused_tokens_account,
            lambda items: f"{len(items)} active API token(s) are owned by non-admin accounts that have been inactive for at least 91 days.",
            "No non-admin active API tokens matched the owner-inactive heuristic.",
            "API token or user inventory was not available to assess unused tokens for accounts.",
            data_available=user_inventory_available and token_inventory_available,
        ),
    ])
    validations.extend(_unsupported_ispm_validations())

    return validations


def _build_validation_summary(validations):
    summary = {
        "assessed": 0,
        "high": 0,
        "medium": 0,
        "low": 0,
        "high_passed": 0,
        "medium_passed": 0,
        "low_passed": 0,
        "high_failed": 0,
        "medium_failed": 0,
        "low_failed": 0,
        "passed": 0,
        "pass_pct": 0,
        "high_pass_pct": 0,
        "medium_pass_pct": 0,
        "low_pass_pct": 0,
    }
    for check in validations or []:
        summary["assessed"] += 1

        severity = str(check.get("severity") or "").strip().lower()
        status = str(check.get("status") or "").strip().lower()
        if severity == "high":
            summary["high"] += 1
            if status == "pass":
                summary["high_passed"] += 1
            if status == "fail":
                summary["high_failed"] += 1
        elif severity in {"moderate", "medium"}:
            summary["medium"] += 1
            if status == "pass":
                summary["medium_passed"] += 1
            if status == "fail":
                summary["medium_failed"] += 1
        elif severity == "low":
            summary["low"] += 1
            if status == "pass":
                summary["low_passed"] += 1
            if status == "fail":
                summary["low_failed"] += 1

        if status == "pass":
            summary["passed"] += 1

    if summary["assessed"]:
        summary["pass_pct"] = round((summary["passed"] / summary["assessed"]) * 100)
    if summary["high"]:
        summary["high_pass_pct"] = round((summary["high_passed"] / summary["high"]) * 100)
    if summary["medium"]:
        summary["medium_pass_pct"] = round((summary["medium_passed"] / summary["medium"]) * 100)
    if summary["low"]:
        summary["low_pass_pct"] = round((summary["low_passed"] / summary["low"]) * 100)

    return summary


def _build_evaluate_summary(sections, domain, extra_context=None):
    section_by_id = _section_map(sections)
    total_sections = len(sections or [])
    populated_sections = sum(1 for s in (sections or []) if s.get("rows"))
    empty_sections = total_sections - populated_sections
    total_rows = sum(len(s.get("rows") or []) for s in (sections or []))

    readiness_groups = [
        {
            "name": "Core Access Controls",
            "section_ids": [
                "groups",
                "group-rules",
                "network-zones",
                "authenticators",
                "password-policies",
                "global-session-policies",
                "authentication-policies",
                "mfa-enrollment-policies",
            ],
        },
        {
            "name": "Identity Federation",
            "section_ids": [
                "identity-providers",
                "idp-discovery-policies",
                "profile-mappings",
            ],
        },
        {
            "name": "Platform Security & Settings",
            "section_ids": [
                "org-settings",
                "security-settings",
                "trusted-origins",
            ],
        },
        {
            "name": "Admin Delegation",
            "section_ids": [
                "custom-admin-roles",
                "resource-sets",
                "admin-assignments-users",
                "admin-assignments-groups",
                "admin-assignments-apps",
            ],
        },
        {
            "name": "Branding & End-User Experience",
            "section_ids": [
                "brand-settings",
                "brand-pages",
                "brand-email-templates",
            ],
        },
    ]

    group_results = []
    score_total = 0
    score_max = 0
    for group in readiness_groups:
        items = []
        present = 0
        for section_id in group["section_ids"]:
            section = section_by_id.get(section_id) or {}
            row_count = len(section.get("rows") or [])
            has_data = row_count > 0
            present += int(has_data)
            items.append(
                {
                    "id": section_id,
                    "title": section.get("title") or section_id,
                    "row_count": row_count,
                    "status": "Present" if has_data else "Not Found / Empty",
                }
            )
        total = len(group["section_ids"])
        pct = round((present / total) * 100) if total else 0
        score_total += present
        score_max += total
        if pct >= 85:
            risk = "Low"
        elif pct >= 60:
            risk = "Medium"
        else:
            risk = "High"
        group_results.append(
            {
                "name": group["name"],
                "present": present,
                "total": total,
                "pct": pct,
                "risk": risk,
                "items": items,
            }
        )

    overall_score = round((score_total / score_max) * 100) if score_max else 0
    if overall_score >= 85:
        readiness_band = "Ready"
    elif overall_score >= 65:
        readiness_band = "Needs Review"
    else:
        readiness_band = "High Attention"

    recommendations = []
    if empty_sections:
        recommendations.append(
            f"{empty_sections} snapshot sections returned no data. Verify API permissions and tenant feature availability."
        )
    if overall_score < 85:
        recommendations.append(
            "Run OktaCompare against the target environment before cutover to confirm policy, app, and IdP parity."
        )
    if not (section_by_id.get("custom-admin-roles") or {}).get("rows"):
        recommendations.append(
            "Review admin delegation model manually if custom admin roles/resource sets are not configured in this tenant."
        )
    if not (section_by_id.get("brand-pages") or {}).get("rows"):
        recommendations.append(
            "Validate end-user branding pages and email templates before go-live if branded experiences are in scope."
        )
    if not recommendations:
        recommendations.append("Tenant configuration coverage looks healthy. Proceed with detailed migration validation.")

    security_validations = _run_security_validations(sections or [], extra_context=extra_context)
    validation_summary = _build_validation_summary(security_validations)

    return {
        "domain": domain,
        "generated_at": datetime.now(ZoneInfo("Australia/Brisbane")).strftime("%Y-%m-%d %H:%M:%S %Z"),
        "total_sections": total_sections,
        "populated_sections": populated_sections,
        "empty_sections": empty_sections,
        "total_rows": total_rows,
        "overall_score": overall_score,
        "readiness_band": readiness_band,
        "groups": group_results,
        "check_prefix_legend": CHECK_PREFIX_LEGEND,
        "security_validations": security_validations,
        "validation_summary": validation_summary,
        "recommendations": recommendations,
    }


def _build_migration_plan(form_data):
    source_domain = form_data.get("source_domain", "").strip()
    target_domain = form_data.get("target_domain", "").strip()
    in_scope = {
        "groups": form_data.get("scope_groups") == "on",
        "apps": form_data.get("scope_apps") == "on",
        "policies": form_data.get("scope_policies") == "on",
        "idp": form_data.get("scope_idp") == "on",
        "branding": form_data.get("scope_branding") == "on",
        "admin": form_data.get("scope_admin") == "on",
    }

    scope_labels = []
    if in_scope["groups"]:
        scope_labels.append("Groups")
    if in_scope["apps"]:
        scope_labels.append("Applications")
    if in_scope["policies"]:
        scope_labels.append("Policies")
    if in_scope["idp"]:
        scope_labels.append("Identity Providers")
    if in_scope["branding"]:
        scope_labels.append("Branding")
    if in_scope["admin"]:
        scope_labels.append("Admin Delegation")

    phases = [
        {
            "phase": "1. Discovery & Baseline",
            "objective": "Capture current-state configuration and confirm migration scope.",
            "tasks": [
                f"Run OktaSnapshot for source tenant ({source_domain}) and export PDF/DOCX baseline.",
                "Document in-scope apps, policies, integrations, and business owners.",
                "Confirm cutover constraints, freeze windows, and rollback expectations.",
            ],
            "outputs": ["Source snapshot baseline", "Scope inventory", "Owner map"],
        },
        {
            "phase": "2. Target Readiness",
            "objective": "Prepare target tenant and validate prerequisite controls.",
            "tasks": [
                f"Run OktaSnapshot for target tenant ({target_domain}) to capture starting state.",
                "Validate org/security settings, trusted origins, authenticator posture, and admin access.",
                "Define target baseline and naming conventions for migrated objects.",
            ],
            "outputs": ["Target readiness checklist", "Target baseline snapshot"],
        },
        {
            "phase": "3. Build / Migration Execution",
            "objective": "Migrate in-scope configuration in controlled waves.",
            "tasks": [
                "Migrate by dependency order (foundational settings -> policies -> apps -> branding).",
                "Track exceptions, unsupported settings, and manual remediation items.",
                "Validate each wave with owners before proceeding.",
            ],
            "outputs": ["Wave tracker", "Issue log", "Validation sign-offs"],
        },
        {
            "phase": "4. Compare & Validate",
            "objective": "Confirm parity and identify residual drift before cutover.",
            "tasks": [
                f"Run OktaCompare between source ({source_domain}) and target ({target_domain}).",
                "Review Critical/Medium findings and remediate before cutover approval.",
                "Export CSV comparison report for approval and audit trail.",
            ],
            "outputs": ["Comparison report", "Remediation list", "Cutover readiness decision"],
        },
        {
            "phase": "5. Cutover & Hypercare",
            "objective": "Execute cutover safely and verify production behavior.",
            "tasks": [
                "Execute cutover checklist during approved maintenance window.",
                "Run post-cutover OktaCompare/OktaSnapshot validation and confirm no unexpected drift.",
                "Track incidents, user impact, and stabilization actions during hypercare.",
            ],
            "outputs": ["Cutover log", "Post-cutover validation", "Hypercare summary"],
        },
    ]

    risks = [
        {
            "name": "Policy/Rule Drift",
            "severity": "High" if in_scope["policies"] else "Medium",
            "mitigation": "Use OktaCompare before cutover and remediate Critical/Medium findings.",
        },
        {
            "name": "App Assignment / Integration Gaps",
            "severity": "High" if in_scope["apps"] else "Low",
            "mitigation": "Validate application owners, assignments, and test sign-on flows by wave.",
        },
        {
            "name": "Federation / IdP Differences",
            "severity": "High" if in_scope["idp"] else "Low",
            "mitigation": "Validate IdP discovery policies, profile mappings, and routing rules in test flows.",
        },
        {
            "name": "Branding / User Experience Regression",
            "severity": "Medium" if in_scope["branding"] else "Low",
            "mitigation": "Snapshot and validate brand pages, email templates, and sign-in UX before go-live.",
        },
        {
            "name": "Admin Access / Operational Readiness",
            "severity": "Medium" if in_scope["admin"] else "Low",
            "mitigation": "Confirm admin assignments, roles, and break-glass access in target tenant.",
        },
    ]

    workflow_name = "Extract Source -> Compare with Target -> Update Missing/Different"

    assumptions = [
        f"In scope: {', '.join(scope_labels) if scope_labels else 'No scopes selected (select at least one for a stronger plan).'}",
    ]

    return {
        "generated_at": datetime.now(ZoneInfo("Australia/Brisbane")).strftime("%Y-%m-%d %H:%M:%S %Z"),
        "source_domain": source_domain,
        "target_domain": target_domain,
        "workflow_name": workflow_name,
        "in_scope": scope_labels,
        "phases": phases,
        "risks": risks,
        "assumptions": assumptions,
    }


def _build_group_sync_summary(source_groups, target_groups):
    source_groups = [g for g in (source_groups or []) if str(g.get("type") or "").upper() == "OKTA_GROUP"]
    target_groups = [g for g in (target_groups or []) if str(g.get("type") or "").upper() == "OKTA_GROUP"]

    src_by_name = {((g.get("profile") or {}).get("name") or ""): g for g in source_groups if (g.get("profile") or {}).get("name")}
    tgt_by_name = {((g.get("profile") or {}).get("name") or ""): g for g in target_groups if (g.get("profile") or {}).get("name")}

    missing = []
    different = []
    matched = 0
    extra = 0
    all_groups = []

    for name, src in src_by_name.items():
        src_profile = src.get("profile") or {}
        src_desc = src_profile.get("description") or ""
        tgt = tgt_by_name.get(name)
        if not tgt:
            missing.append({
                "name": name,
                "description": src_desc,
            })
            all_groups.append({
                "name": name,
                "source_description": src_desc,
                "target_description": "",
                "status": "Missing in Target",
            })
            continue
        tgt_profile = tgt.get("profile") or {}
        tgt_desc = tgt_profile.get("description") or ""
        if (src_desc or "") != (tgt_desc or ""):
            different.append({
                "name": name,
                "source_description": src_desc,
                "target_description": tgt_desc,
                "target_id": tgt.get("id"),
            })
            all_groups.append({
                "name": name,
                "source_description": src_desc,
                "target_description": tgt_desc,
                "status": "Different",
            })
        else:
            matched += 1
            all_groups.append({
                "name": name,
                "source_description": src_desc,
                "target_description": tgt_desc,
                "status": "Match",
            })

    for name in tgt_by_name:
        if name not in src_by_name:
            extra += 1
            tgt_desc = ((tgt_by_name.get(name) or {}).get("profile") or {}).get("description") or ""
            all_groups.append({
                "name": name,
                "source_description": "",
                "target_description": tgt_desc,
                "status": "Extra in Target",
            })

    all_groups.sort(key=lambda g: ((g.get("name") or "").lower(), g.get("status") or ""))

    return {
        "source_count": len(src_by_name),
        "target_count": len(tgt_by_name),
        "missing": missing,
        "different": different,
        "all_groups": all_groups,
        "matched_count": matched,
        "extra_count": extra,
        "pending_count": len(missing) + len(different),
    }


def _okta_headers(api_token):
    return {
        "Authorization": f"SSWS {api_token}",
        "Accept": "application/json",
        "Content-Type": "application/json",
    }


def _create_group(domain, api_token, name, description=""):
    url = f"{_ensure_https_domain(domain).rstrip('/')}/api/v1/groups"
    payload = {"profile": {"name": name, "description": description or ""}}
    return requests.post(url, headers=_okta_headers(api_token), json=payload, timeout=30)


def _update_group_description(domain, api_token, group_id, name, description=""):
    url = f"{_ensure_https_domain(domain).rstrip('/')}/api/v1/groups/{group_id}"
    payload = {"profile": {"name": name, "description": description or ""}}
    return requests.put(url, headers=_okta_headers(api_token), json=payload, timeout=30)


def _oktaevaluate_csv_bytes(evaluation):
    output = io.StringIO()
    writer = csv.DictWriter(
        output,
        fieldnames=["Check ID", "What Was Checked", "Result", "Severity", "Summary", "Details"],
    )
    writer.writeheader()
    for check in (evaluation or {}).get("security_validations", []) or []:
        writer.writerow(
            {
                "Check ID": check.get("check_id", ""),
                "What Was Checked": check.get("title", ""),
                "Result": check.get("status", ""),
                "Severity": check.get("severity", ""),
                "Summary": check.get("summary", ""),
                "Details": " | ".join([str(i) for i in (check.get("items") or [])]),
            }
        )
    return output.getvalue().encode("utf-8")


# ---------------------------------------------------
# Compare Groups (unchanged logic)
# ---------------------------------------------------
def compare_groups(groupsA, groupsB):
    diffs = []
    matches = []

    dictA = {g["profile"]["name"]: g for g in groupsA}
    dictB = {g["profile"]["name"]: g for g in groupsB}

    for name, grpA in dictA.items():
        descA = grpA["profile"].get("description", "")

        if name not in dictB:
            diffs.append({
                "Category": "Groups",
                "Object": name,
                "Attribute": "-",
                "Env A Value": "Exists",
                "Env B Value": "Missing",
                "Difference Type": "Missing in Env B",
                "Impact": "Access",
                "Recommended Action": f"Create group '{name}' in Env B",
                "Priority": "🔴 Critical"
            })
        else:
            descB = dictB[name]["profile"].get("description", "")
            if descA != descB:
                diffs.append({
                    "Category": "Groups",
                    "Object": name,
                    "Attribute": "Description",
                    "Env A Value": descA,
                    "Env B Value": descB,
                    "Difference Type": "Mismatch",
                    "Impact": "Configuration Drift",
                    "Recommended Action": f"Align description for group '{name}'",
                    "Priority": "🟠 Medium"
                })
            else:
                matches.append({
                    "Category": "Groups",
                    "Object": name,
                    "Attribute": "Description",
                    "Value": descA
                })

    # Extra in Env B
    for name in dictB:
        if name not in dictA:
            diffs.append({
                "Category": "Groups",
                "Object": name,
                "Attribute": "-",
                "Env A Value": "Missing",
                "Env B Value": "Exists",
                "Difference Type": "Extra in Env B",
                "Impact": "Configuration Drift",
                "Recommended Action": f"Review/remove group '{name}' in Env B",
                "Priority": "🟡 Low"
            })

    return diffs, matches



# ---------------------------------------------------
# Main Page
# ---------------------------------------------------
@app.route("/", methods=["GET", "POST"])
def index():
    logger.info("Index request received: method=%s", request.method)
    if request.method == "POST":
        session.clear()

        # -------------
        # Inputs
        # -------------
        logger.info("Collecting input values.")
        envA_domain = request.form.get("envA_domain", "").strip()
        envA_token  = request.form.get("envA_token", "").strip()
        envB_domain = request.form.get("envB_domain", "").strip()
        envB_token  = request.form.get("envB_token", "").strip()

        if not all([envA_domain, envA_token, envB_domain, envB_token]):
            logger.warning("Missing required comparison inputs.")
            return render_template(
                "oktacompare_error.html",
                title="Missing Required Input",
                message="Please provide Env A and Env B domains and API tokens.",
            ), 400


        # ===================================================
        # GROUPS
        # ===================================================
        logger.info("Fetching groups for envA and envB.")
        groupsA = get_groups(envA_domain, envA_token)
        groupsB = get_groups(envB_domain, envB_token)

        logger.info("Comparing groups.")
        group_diffs, group_matches_raw = compare_groups(groupsA, groupsB)

        group_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            }
            for m in group_matches_raw
        ]

        group_df = pd.DataFrame(group_diffs + group_matches_display)
        group_summary_counts = pd.DataFrame(group_diffs)["Priority"].value_counts().to_dict() if group_diffs else {}
        group_total_diff = len(group_diffs)
        logger.info("Groups comparison complete: diffs=%s matches=%s", len(group_diffs), len(group_matches_raw))


        # ===================================================
        # GROUP RULES
        # ===================================================
        logger.info("Comparing group rules.")
        rule_diffs, rule_matches_raw = compare_group_rules(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        rule_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in rule_matches_raw
        ]

        rule_df = pd.DataFrame(rule_diffs + rule_matches_display)
        rule_summary_counts = pd.DataFrame(rule_diffs)["Priority"].value_counts().to_dict() if rule_diffs else {}
        rule_total_diff = len(rule_diffs)
        logger.info("Group rules comparison complete: diffs=%s matches=%s", len(rule_diffs), len(rule_matches_raw))


        # ===================================================
        # NETWORK ZONES
        # ===================================================
        logger.info("Comparing network zones.")
        zone_diffs, zone_matches_raw = compare_network_zones(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        zone_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m.get("Value", ""),
                "Env B Value": m.get("Value", ""),
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in zone_matches_raw
        ]

        zone_df = pd.DataFrame(zone_diffs + zone_matches_display)
        zone_summary_counts = pd.DataFrame(zone_diffs)["Priority"].value_counts().to_dict() if zone_diffs else {}
        zone_total_diff = len(zone_diffs)
        logger.info("Network zones comparison complete: diffs=%s matches=%s", len(zone_diffs), len(zone_matches_raw))


        # ===================================================
        # APPLICATIONS
        # ===================================================
        logger.info("Comparing applications.")
        app_diffs, app_matches_raw = compare_applications(
            envA_domain, envA_token,
            envB_domain, envB_token,
            compare_group_assignments=False,
        )

        app_matches_display = [
            {
                "Category": m.get("Category", "Applications"),
                "Object": m.get("Object"),
                "Attribute": m.get("Attribute"),
                "Env A Value": m.get("Value", ""),
                "Env B Value": m.get("Value", ""),
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in app_matches_raw
        ]

        apps_df = pd.DataFrame(app_diffs + app_matches_display)
        app_summary_counts = pd.DataFrame(app_diffs)["Priority"].value_counts().to_dict() if app_diffs else {}
        app_total_diff = len(app_diffs)
        logger.info("Applications comparison complete: diffs=%s matches=%s", len(app_diffs), len(app_matches_raw))

        # ===================================================
        # AUTHENTICATORS
        # ===================================================
        logger.info("Comparing authenticators.")
        auth_diffs, auth_matches_raw = compare_authenticators(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        auth_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in auth_matches_raw
        ]

        auth_df = pd.DataFrame(auth_diffs + auth_matches_display)
        auth_summary_counts = pd.DataFrame(auth_diffs)["Priority"].value_counts().to_dict() if auth_diffs else {}
        auth_total_diff = len(auth_diffs)
        logger.info("Authenticators comparison complete: diffs=%s matches=%s", len(auth_diffs), len(auth_matches_raw))

        # ===================================================
        # AUTHENTICATOR ENROLLMENT POLICIES (MFA)
        # ===================================================
        logger.info("Comparing authenticator enrollment policies.")
        mfa_diffs, mfa_matches_raw = compare_mfa_policies(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        mfa_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in mfa_matches_raw
        ]

        mfa_df = pd.DataFrame(mfa_diffs + mfa_matches_display)
        mfa_summary_counts = pd.DataFrame(mfa_diffs)["Priority"].value_counts().to_dict() if mfa_diffs else {}
        mfa_total_diff = len(mfa_diffs)
        logger.info("MFA policies comparison complete: diffs=%s matches=%s", len(mfa_diffs), len(mfa_matches_raw))

        # ===================================================
        # PASSWORD POLICIES
        # ===================================================
        logger.info("Comparing password policies.")
        pwd_diffs, pwd_matches_raw = compare_password_policies(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        pwd_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in pwd_matches_raw
        ]

        pwd_df = pd.DataFrame(pwd_diffs + pwd_matches_display)
        pwd_summary_counts = pd.DataFrame(pwd_diffs)["Priority"].value_counts().to_dict() if pwd_diffs else {}
        pwd_total_diff = len(pwd_diffs)
        logger.info("Password policies comparison complete: diffs=%s matches=%s", len(pwd_diffs), len(pwd_matches_raw))

        # ===================================================
        # APP SIGN-ON POLICIES (ACCESS_POLICY)
        # ===================================================
        logger.info("Comparing app sign-on policies.")
        access_diffs, access_matches_raw = compare_access_policies(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        access_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in access_matches_raw
        ]

        access_df = pd.DataFrame(access_diffs + access_matches_display)
        access_summary_counts = pd.DataFrame(access_diffs)["Priority"].value_counts().to_dict() if access_diffs else {}
        access_total_diff = len(access_diffs)
        logger.info("App sign-on policies comparison complete: diffs=%s matches=%s", len(access_diffs), len(access_matches_raw))

        # ===================================================
        # IDP DISCOVERY POLICIES
        # ===================================================
        logger.info("Comparing IDP discovery policies.")
        idp_diffs, idp_matches_raw = compare_idp_discovery_policies(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        idp_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in idp_matches_raw
        ]

        idp_df = pd.DataFrame(idp_diffs + idp_matches_display)
        idp_summary_counts = pd.DataFrame(idp_diffs)["Priority"].value_counts().to_dict() if idp_diffs else {}
        idp_total_diff = len(idp_diffs)
        logger.info("IDP discovery policies comparison complete: diffs=%s matches=%s", len(idp_diffs), len(idp_matches_raw))

        # ===================================================
        # PROFILE ENROLLMENT POLICIES
        # ===================================================
        logger.info("Comparing profile enrollment policies.")
        profile_diffs, profile_matches_raw = compare_profile_enrollment_policies(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        profile_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in profile_matches_raw
        ]

        profile_df = pd.DataFrame(profile_diffs + profile_matches_display)
        profile_summary_counts = pd.DataFrame(profile_diffs)["Priority"].value_counts().to_dict() if profile_diffs else {}
        profile_total_diff = len(profile_diffs)
        logger.info(
            "Profile enrollment policies comparison complete: diffs=%s matches=%s",
            len(profile_diffs),
            len(profile_matches_raw),
        )

        # ===================================================
        # ENTITY RISK POLICIES
        # ===================================================
        logger.info("Comparing entity risk policies.")
        entity_risk_diffs, entity_risk_matches_raw = compare_entity_risk_policies(
            envA_domain, envA_token,
            envB_domain, envB_token
        )
        entity_risk_matches_display = [
            {
                "Category": m["Category"], "Object": m["Object"], "Attribute": m["Attribute"],
                "Env A Value": m["Value"], "Env B Value": m["Value"], "Difference Type": "Match",
                "Impact": "", "Recommended Action": "", "Priority": "🟢 Match"
            } for m in entity_risk_matches_raw
        ]
        entity_risk_df = pd.DataFrame(entity_risk_diffs + entity_risk_matches_display)
        entity_risk_summary_counts = pd.DataFrame(entity_risk_diffs)["Priority"].value_counts().to_dict() if entity_risk_diffs else {}
        entity_risk_total_diff = len(entity_risk_diffs)
        logger.info("Entity risk policies comparison complete: diffs=%s matches=%s", len(entity_risk_diffs), len(entity_risk_matches_raw))

        # ===================================================
        # IDENTITY THREAT PROTECTION POLICIES
        # ===================================================
        logger.info("Comparing identity threat protection policies.")
        post_auth_diffs, post_auth_matches_raw = compare_post_auth_session_policies(
            envA_domain, envA_token,
            envB_domain, envB_token
        )
        post_auth_matches_display = [
            {
                "Category": m["Category"], "Object": m["Object"], "Attribute": m["Attribute"],
                "Env A Value": m["Value"], "Env B Value": m["Value"], "Difference Type": "Match",
                "Impact": "", "Recommended Action": "", "Priority": "🟢 Match"
            } for m in post_auth_matches_raw
        ]
        post_auth_df = pd.DataFrame(post_auth_diffs + post_auth_matches_display)
        post_auth_summary_counts = pd.DataFrame(post_auth_diffs)["Priority"].value_counts().to_dict() if post_auth_diffs else {}
        post_auth_total_diff = len(post_auth_diffs)
        logger.info("Identity threat protection policies comparison complete: diffs=%s matches=%s", len(post_auth_diffs), len(post_auth_matches_raw))

        # ===================================================
        # BRAND SETTINGS
        # ===================================================
        logger.info("Comparing brand settings.")
        brand_diffs, brand_matches_raw = compare_brand_settings(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        brand_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in brand_matches_raw
        ]

        brand_df = pd.DataFrame(brand_diffs + brand_matches_display)
        brand_summary_counts = pd.DataFrame(brand_diffs)["Priority"].value_counts().to_dict() if brand_diffs else {}
        brand_total_diff = len(brand_diffs)
        logger.info("Brand settings comparison complete: diffs=%s matches=%s", len(brand_diffs), len(brand_matches_raw))

        # ===================================================
        # BRAND PAGES
        # ===================================================
        logger.info("Comparing brand pages.")
        brand_pages_diffs, brand_pages_matches_raw = compare_brand_pages(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        brand_pages_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in brand_pages_matches_raw
        ]

        brand_pages_df = pd.DataFrame(brand_pages_diffs + brand_pages_matches_display)
        brand_pages_summary_counts = (
            pd.DataFrame(brand_pages_diffs)["Priority"].value_counts().to_dict() if brand_pages_diffs else {}
        )
        brand_pages_total_diff = len(brand_pages_diffs)
        logger.info(
            "Brand pages comparison complete: diffs=%s matches=%s",
            len(brand_pages_diffs),
            len(brand_pages_matches_raw),
        )

        # ===================================================
        # BRAND EMAIL TEMPLATES
        # ===================================================
        logger.info("Comparing brand email templates.")
        brand_email_diffs, brand_email_matches_raw = compare_brand_email_templates(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        brand_email_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in brand_email_matches_raw
        ]

        brand_email_df = pd.DataFrame(brand_email_diffs + brand_email_matches_display)
        brand_email_summary_counts = (
            pd.DataFrame(brand_email_diffs)["Priority"].value_counts().to_dict() if brand_email_diffs else {}
        )
        brand_email_total_diff = len(brand_email_diffs)
        logger.info(
            "Brand email templates comparison complete: diffs=%s matches=%s",
            len(brand_email_diffs),
            len(brand_email_matches_raw),
        )

        # ===================================================
        # AUTHORIZATION SERVERS - SETTINGS
        # ===================================================
        logger.info("Comparing authorization servers settings.")
        authz_diffs, authz_matches_raw = compare_authorization_servers_settings(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        authz_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in authz_matches_raw
        ]

        authz_df = pd.DataFrame(authz_diffs + authz_matches_display)
        authz_summary_counts = pd.DataFrame(authz_diffs)["Priority"].value_counts().to_dict() if authz_diffs else {}
        authz_total_diff = len(authz_diffs)
        logger.info("Authorization servers settings comparison complete: diffs=%s matches=%s", len(authz_diffs), len(authz_matches_raw))

        # ===================================================
        # AUTHORIZATION SERVERS - ACCESS POLICIES
        # ===================================================
        logger.info("Comparing authorization servers access policies.")
        authz_policy_diffs, authz_policy_matches_raw = compare_authorization_servers_access_policies(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        authz_policy_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in authz_policy_matches_raw
        ]

        authz_policy_df = pd.DataFrame(authz_policy_diffs + authz_policy_matches_display)
        authz_policy_summary_counts = (
            pd.DataFrame(authz_policy_diffs)["Priority"].value_counts().to_dict() if authz_policy_diffs else {}
        )
        authz_policy_total_diff = len(authz_policy_diffs)
        logger.info(
            "Authorization servers access policies comparison complete: diffs=%s matches=%s",
            len(authz_policy_diffs),
            len(authz_policy_matches_raw),
        )

        # ===================================================
        # CUSTOM ADMIN ROLES
        # ===================================================
        logger.info("Comparing custom admin roles.")
        admin_role_diffs, admin_role_matches_raw = compare_custom_admin_roles(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        admin_role_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in admin_role_matches_raw
        ]

        admin_role_df = pd.DataFrame(admin_role_diffs + admin_role_matches_display)
        admin_role_summary_counts = (
            pd.DataFrame(admin_role_diffs)["Priority"].value_counts().to_dict() if admin_role_diffs else {}
        )
        admin_role_total_diff = len(admin_role_diffs)
        logger.info(
            "Custom admin roles comparison complete: diffs=%s matches=%s",
            len(admin_role_diffs),
            len(admin_role_matches_raw),
        )

        # ===================================================
        # RESOURCE SETS
        # ===================================================
        logger.info("Comparing resource sets.")
        resource_set_diffs, resource_set_matches_raw = compare_resource_sets(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        resource_set_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in resource_set_matches_raw
        ]

        resource_set_df = pd.DataFrame(resource_set_diffs + resource_set_matches_display)
        resource_set_summary_counts = (
            pd.DataFrame(resource_set_diffs)["Priority"].value_counts().to_dict() if resource_set_diffs else {}
        )
        resource_set_total_diff = len(resource_set_diffs)
        logger.info(
            "Resource sets comparison complete: diffs=%s matches=%s",
            len(resource_set_diffs),
            len(resource_set_matches_raw),
        )

        # ===================================================
        # ADMIN ASSIGNMENTS
        # ===================================================
        logger.info("Comparing admin assignments.")
        admin_assign_diffs, admin_assign_matches_raw = compare_admin_assignments(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        admin_assign_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in admin_assign_matches_raw
        ]

        admin_assign_df = pd.DataFrame(admin_assign_diffs + admin_assign_matches_display)
        admin_assign_summary_counts = (
            pd.DataFrame(admin_assign_diffs)["Priority"].value_counts().to_dict() if admin_assign_diffs else {}
        )
        admin_assign_total_diff = len(admin_assign_diffs)
        logger.info(
            "Admin assignments comparison complete: diffs=%s matches=%s",
            len(admin_assign_diffs),
            len(admin_assign_matches_raw),
        )

        # ===================================================
        # API TOKENS
        # ===================================================
        logger.info("Comparing API tokens.")
        api_token_diffs, api_token_matches_raw = compare_api_tokens(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        api_token_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in api_token_matches_raw
        ]

        api_token_df = pd.DataFrame(api_token_diffs + api_token_matches_display)
        api_token_summary_counts = (
            pd.DataFrame(api_token_diffs)["Priority"].value_counts().to_dict() if api_token_diffs else {}
        )
        api_token_total_diff = len(api_token_diffs)
        logger.info(
            "API tokens comparison complete: diffs=%s matches=%s",
            len(api_token_diffs),
            len(api_token_matches_raw),
        )

        # ===================================================
        # SECURITY GENERAL SETTINGS
        # ===================================================
        logger.info("Comparing security general settings.")
        sec_diffs, sec_matches_raw = compare_security_general_settings(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        sec_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in sec_matches_raw
        ]

        sec_df = pd.DataFrame(sec_diffs + sec_matches_display)
        sec_summary_counts = (
            pd.DataFrame(sec_diffs)["Priority"].value_counts().to_dict() if sec_diffs else {}
        )
        sec_total_diff = len(sec_diffs)
        logger.info(
            "Security general settings comparison complete: diffs=%s matches=%s",
            len(sec_diffs),
            len(sec_matches_raw),
        )

        # ===================================================
        # ORG GENERAL SETTINGS
        # ===================================================
        logger.info("Comparing org general settings.")
        org_diffs, org_matches_raw = compare_org_settings(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        org_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in org_matches_raw
        ]

        org_df = pd.DataFrame(org_diffs + org_matches_display)
        org_summary_counts = (
            pd.DataFrame(org_diffs)["Priority"].value_counts().to_dict() if org_diffs else {}
        )
        org_total_diff = len(org_diffs)
        logger.info(
            "Org general settings comparison complete: diffs=%s matches=%s",
            len(org_diffs),
            len(org_matches_raw),
        )

        # ===================================================
        # IDENTITY PROVIDERS
        # ===================================================
        logger.info("Comparing identity providers.")
        idp_provider_diffs, idp_provider_matches_raw = compare_identity_providers(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        idp_provider_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in idp_provider_matches_raw
        ]

        idp_provider_df = pd.DataFrame(idp_provider_diffs + idp_provider_matches_display)
        idp_provider_summary_counts = (
            pd.DataFrame(idp_provider_diffs)["Priority"].value_counts().to_dict() if idp_provider_diffs else {}
        )
        idp_provider_total_diff = len(idp_provider_diffs)
        logger.info(
            "Identity providers comparison complete: diffs=%s matches=%s",
            len(idp_provider_diffs),
            len(idp_provider_matches_raw),
        )

        # ===================================================
        # REALMS
        # ===================================================
        logger.info("Comparing realms.")
        realm_diffs, realm_matches_raw = compare_realms(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        realm_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in realm_matches_raw
        ]

        realm_df = pd.DataFrame(realm_diffs + realm_matches_display)
        realm_summary_counts = (
            pd.DataFrame(realm_diffs)["Priority"].value_counts().to_dict() if realm_diffs else {}
        )
        realm_total_diff = len(realm_diffs)
        logger.info(
            "Realms comparison complete: diffs=%s matches=%s",
            len(realm_diffs),
            len(realm_matches_raw),
        )

        # ===================================================
        # REALM ASSIGNMENTS
        # ===================================================
        logger.info("Comparing realm assignments.")
        realm_assign_diffs, realm_assign_matches_raw = compare_realm_assignments(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        realm_assign_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in realm_assign_matches_raw
        ]

        realm_assign_df = pd.DataFrame(realm_assign_diffs + realm_assign_matches_display)
        realm_assign_summary_counts = (
            pd.DataFrame(realm_assign_diffs)["Priority"].value_counts().to_dict() if realm_assign_diffs else {}
        )
        realm_assign_total_diff = len(realm_assign_diffs)
        logger.info(
            "Realm assignments comparison complete: diffs=%s matches=%s",
            len(realm_assign_diffs),
            len(realm_assign_matches_raw),
        )

        # ===================================================
        # PROFILE SCHEMA - USER
        # ===================================================
        logger.info("Comparing user profile schema.")
        schema_diffs, schema_matches_raw = compare_user_profile_schema(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        schema_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in schema_matches_raw
        ]

        schema_df = pd.DataFrame(schema_diffs + schema_matches_display)
        schema_summary_counts = (
            pd.DataFrame(schema_diffs)["Priority"].value_counts().to_dict() if schema_diffs else {}
        )
        schema_total_diff = len(schema_diffs)
        logger.info(
            "User profile schema comparison complete: diffs=%s matches=%s",
            len(schema_diffs),
            len(schema_matches_raw),
        )

        # ===================================================
        # PROFILE MAPPINGS
        # ===================================================
        logger.info("Comparing profile mappings.")
        mapping_diffs, mapping_matches_raw = compare_profile_mappings(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        mapping_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in mapping_matches_raw
        ]

        mapping_df = pd.DataFrame(mapping_diffs + mapping_matches_display)
        mapping_summary_counts = (
            pd.DataFrame(mapping_diffs)["Priority"].value_counts().to_dict() if mapping_diffs else {}
        )
        mapping_total_diff = len(mapping_diffs)
        logger.info(
            "Profile mappings comparison complete: diffs=%s matches=%s",
            len(mapping_diffs),
            len(mapping_matches_raw),
        )

        # ===================================================
        # TRUSTED ORIGINS
        # ===================================================
        logger.info("Comparing trusted origins.")
        origin_diffs, origin_matches_raw = compare_trusted_origins(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        origin_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in origin_matches_raw
        ]

        origin_df = pd.DataFrame(origin_diffs + origin_matches_display)
        origin_summary_counts = (
            pd.DataFrame(origin_diffs)["Priority"].value_counts().to_dict() if origin_diffs else {}
        )
        origin_total_diff = len(origin_diffs)
        logger.info(
            "Trusted origins comparison complete: diffs=%s matches=%s",
            len(origin_diffs),
            len(origin_matches_raw),
        )

        # ===================================================
        # EVENT HOOKS
        # ===================================================
        logger.info("Comparing event hooks.")
        event_hook_diffs, event_hook_matches_raw = compare_event_hooks(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        event_hook_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in event_hook_matches_raw
        ]

        event_hook_df = pd.DataFrame(event_hook_diffs + event_hook_matches_display)
        event_hook_summary_counts = (
            pd.DataFrame(event_hook_diffs)["Priority"].value_counts().to_dict() if event_hook_diffs else {}
        )
        event_hook_total_diff = len(event_hook_diffs)
        logger.info(
            "Event hooks comparison complete: diffs=%s matches=%s",
            len(event_hook_diffs),
            len(event_hook_matches_raw),
        )

        # ===================================================
        # INLINE HOOKS
        # ===================================================
        logger.info("Comparing inline hooks.")
        inline_hook_diffs, inline_hook_matches_raw = compare_inline_hooks(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        inline_hook_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in inline_hook_matches_raw
        ]

        inline_hook_df = pd.DataFrame(inline_hook_diffs + inline_hook_matches_display)
        inline_hook_summary_counts = (
            pd.DataFrame(inline_hook_diffs)["Priority"].value_counts().to_dict() if inline_hook_diffs else {}
        )
        inline_hook_total_diff = len(inline_hook_diffs)
        logger.info(
            "Inline hooks comparison complete: diffs=%s matches=%s",
            len(inline_hook_diffs),
            len(inline_hook_matches_raw),
        )

        # ===================================================
        # ACCESS CONTROLS - ATTACK PROTECTION
        # ===================================================
        logger.info("Comparing attack protection controls.")
        attack_protection_diffs, attack_protection_matches_raw = compare_attack_protection(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        attack_protection_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in attack_protection_matches_raw
        ]

        attack_protection_df = pd.DataFrame(attack_protection_diffs + attack_protection_matches_display)
        attack_protection_summary_counts = (
            pd.DataFrame(attack_protection_diffs)["Priority"].value_counts().to_dict() if attack_protection_diffs else {}
        )
        attack_protection_total_diff = len(attack_protection_diffs)
        logger.info(
            "Attack protection comparison complete: diffs=%s matches=%s",
            len(attack_protection_diffs),
            len(attack_protection_matches_raw),
        )

        # ===================================================
        # GROUP PUSH MAPPINGS
        # ===================================================
        logger.info("Comparing group push mappings.")
        group_push_diffs, group_push_matches_raw = compare_group_push_mappings(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        group_push_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in group_push_matches_raw
        ]

        group_push_df = pd.DataFrame(group_push_diffs + group_push_matches_display)
        group_push_summary_counts = (
            pd.DataFrame(group_push_diffs)["Priority"].value_counts().to_dict() if group_push_diffs else {}
        )
        group_push_total_diff = len(group_push_diffs)
        logger.info(
            "Group push mappings comparison complete: diffs=%s matches=%s",
            len(group_push_diffs),
            len(group_push_matches_raw),
        )

        # ===================================================
        # AGENTS
        # ===================================================
        logger.info("Comparing agents.")
        agent_diffs, agent_matches_raw = compare_agents(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        agent_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            } for m in agent_matches_raw
        ]

        agent_df = pd.DataFrame(agent_diffs + agent_matches_display)
        agent_summary_counts = (
            pd.DataFrame(agent_diffs)["Priority"].value_counts().to_dict() if agent_diffs else {}
        )
        agent_total_diff = len(agent_diffs)
        logger.info(
            "Agents comparison complete: diffs=%s matches=%s",
            len(agent_diffs),
            len(agent_matches_raw),
        )


        # ===================================================
        # GLOBAL SESSION POLICIES (policies + rules together)
        # ===================================================
        logger.info("Comparing session policies.")
        session_diffs, session_matches_raw = compare_session_policies(
            envA_domain, envA_token,
            envB_domain, envB_token
        )

        session_matches_display = [
            {
                "Category": m["Category"],
                "Object": m["Object"],
                "Attribute": m["Attribute"],
                "Env A Value": m["Value"],
                "Env B Value": m["Value"],
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match"
            }
            for m in session_matches_raw
        ]

        session_df = pd.DataFrame(session_diffs + session_matches_display)
        session_summary_counts = pd.DataFrame(session_diffs)["Priority"].value_counts().to_dict() if session_diffs else {}
        session_total_diff = len(session_diffs)
        logger.info("Session policies comparison complete: diffs=%s matches=%s", len(session_diffs), len(session_matches_raw))


        # ===================================================
        # SESSION STORAGE
        # ===================================================
        logger.info("Storing session results.")
        all_diffs = (
            group_diffs
            + rule_diffs
            + zone_diffs
            + app_diffs
            + session_diffs
            + auth_diffs
            + mfa_diffs
            + pwd_diffs
            + access_diffs
            + idp_diffs
            + profile_diffs
            + entity_risk_diffs
            + post_auth_diffs
            + brand_diffs
            + brand_pages_diffs
            + brand_email_diffs
            + authz_diffs
            + authz_policy_diffs
            + admin_role_diffs
            + resource_set_diffs
            + admin_assign_diffs
            + api_token_diffs
            + sec_diffs
            + org_diffs
            + idp_provider_diffs
            + realm_diffs
            + realm_assign_diffs
            + schema_diffs
            + mapping_diffs
            + origin_diffs
            + event_hook_diffs
            + inline_hook_diffs
            + attack_protection_diffs
            + group_push_diffs
            + agent_diffs
        )
        all_matches_raw = (
            group_matches_raw
            + rule_matches_raw
            + zone_matches_raw
            + app_matches_raw
            + session_matches_raw
            + auth_matches_raw
            + mfa_matches_raw
            + pwd_matches_raw
            + access_matches_raw
            + idp_matches_raw
            + profile_matches_raw
            + entity_risk_matches_raw
            + post_auth_matches_raw
            + brand_matches_raw
            + brand_pages_matches_raw
            + brand_email_matches_raw
            + authz_matches_raw
            + authz_policy_matches_raw
            + admin_role_matches_raw
            + resource_set_matches_raw
            + admin_assign_matches_raw
            + api_token_matches_raw
            + sec_matches_raw
            + org_matches_raw
            + idp_provider_matches_raw
            + realm_matches_raw
            + realm_assign_matches_raw
            + schema_matches_raw
            + mapping_matches_raw
            + origin_matches_raw
            + event_hook_matches_raw
            + inline_hook_matches_raw
            + attack_protection_matches_raw
            + group_push_matches_raw
            + agent_matches_raw
        )
        LAST_EXPORT["diffs"] = all_diffs
        LAST_EXPORT["matches"] = all_matches_raw
        export_bytes = (
            len(json.dumps(all_diffs, default=str).encode("utf-8"))
            + len(json.dumps(all_matches_raw, default=str).encode("utf-8"))
        )
        logger.info("Export payload size: %.2f KB", export_bytes / 1024)


        # ===================================================
        # Render Report
        # ===================================================
        logger.info("Rendering report for envA=%s envB=%s.", envA_domain, envB_domain)
        return render_template(
            "oktacompare_report.html",

            # Groups
            group_df=group_df.to_dict(orient="records"),
            group_summary_counts=group_summary_counts,
            group_total_diff=group_total_diff,

            # Group Rules
            rule_df=rule_df.to_dict(orient="records"),
            rule_summary_counts=rule_summary_counts,
            rule_total_diff=rule_total_diff,

            # Network Zones
            zone_df=zone_df.to_dict(orient="records"),
            zone_summary_counts=zone_summary_counts,
            zone_total_diff=zone_total_diff,

            # Applications
            apps_df=apps_df.to_dict(orient="records"),
            app_summary_counts=app_summary_counts,
            app_total_diff=app_total_diff,

            # Authenticators
            auth_df=auth_df.to_dict(orient="records"),
            auth_summary_counts=auth_summary_counts,
            auth_total_diff=auth_total_diff,

            # Authenticator Enrollment Policies (MFA)
            mfa_df=mfa_df.to_dict(orient="records"),
            mfa_summary_counts=mfa_summary_counts,
            mfa_total_diff=mfa_total_diff,

            # Password Policies
            pwd_df=pwd_df.to_dict(orient="records"),
            pwd_summary_counts=pwd_summary_counts,
            pwd_total_diff=pwd_total_diff,

            # App Sign-On Policies
            access_df=access_df.to_dict(orient="records"),
            access_summary_counts=access_summary_counts,
            access_total_diff=access_total_diff,

            # IDP Discovery Policies
            idp_df=idp_df.to_dict(orient="records"),
            idp_summary_counts=idp_summary_counts,
            idp_total_diff=idp_total_diff,

            # Profile Enrollment Policies
            profile_df=profile_df.to_dict(orient="records"),
            profile_summary_counts=profile_summary_counts,
            profile_total_diff=profile_total_diff,

            # Entity Risk Policies
            entity_risk_df=entity_risk_df.to_dict(orient="records"),
            entity_risk_summary_counts=entity_risk_summary_counts,
            entity_risk_total_diff=entity_risk_total_diff,

            # Identity Threat Protection Policies
            post_auth_df=post_auth_df.to_dict(orient="records"),
            post_auth_summary_counts=post_auth_summary_counts,
            post_auth_total_diff=post_auth_total_diff,

            # Brand Settings
            brand_df=brand_df.to_dict(orient="records"),
            brand_summary_counts=brand_summary_counts,
            brand_total_diff=brand_total_diff,

            # Brand Pages
            brand_pages_df=brand_pages_df.to_dict(orient="records"),
            brand_pages_summary_counts=brand_pages_summary_counts,
            brand_pages_total_diff=brand_pages_total_diff,

            # Brand Email Templates
            brand_email_df=brand_email_df.to_dict(orient="records"),
            brand_email_summary_counts=brand_email_summary_counts,
            brand_email_total_diff=brand_email_total_diff,

            # Authorization Servers - Settings
            authz_df=authz_df.to_dict(orient="records"),
            authz_summary_counts=authz_summary_counts,
            authz_total_diff=authz_total_diff,

            # Authorization Servers - Access Policies
            authz_policy_df=authz_policy_df.to_dict(orient="records"),
            authz_policy_summary_counts=authz_policy_summary_counts,
            authz_policy_total_diff=authz_policy_total_diff,

            # Custom Admin Roles
            admin_role_df=admin_role_df.to_dict(orient="records"),
            admin_role_summary_counts=admin_role_summary_counts,
            admin_role_total_diff=admin_role_total_diff,

            # Resource Sets
            resource_set_df=resource_set_df.to_dict(orient="records"),
            resource_set_summary_counts=resource_set_summary_counts,
            resource_set_total_diff=resource_set_total_diff,

            # Admin Assignments
            admin_assign_df=admin_assign_df.to_dict(orient="records"),
            admin_assign_summary_counts=admin_assign_summary_counts,
            admin_assign_total_diff=admin_assign_total_diff,

            # API Tokens
            api_token_df=api_token_df.to_dict(orient="records"),
            api_token_summary_counts=api_token_summary_counts,
            api_token_total_diff=api_token_total_diff,

            # Security General Settings
            sec_df=sec_df.to_dict(orient="records"),
            sec_summary_counts=sec_summary_counts,
            sec_total_diff=sec_total_diff,

            # Org General Settings
            org_df=org_df.to_dict(orient="records"),
            org_summary_counts=org_summary_counts,
            org_total_diff=org_total_diff,

            # Identity Providers
            idp_provider_df=idp_provider_df.to_dict(orient="records"),
            idp_provider_summary_counts=idp_provider_summary_counts,
            idp_provider_total_diff=idp_provider_total_diff,

            # Realms
            realm_df=realm_df.to_dict(orient="records"),
            realm_summary_counts=realm_summary_counts,
            realm_total_diff=realm_total_diff,

            # Realm Assignments
            realm_assign_df=realm_assign_df.to_dict(orient="records"),
            realm_assign_summary_counts=realm_assign_summary_counts,
            realm_assign_total_diff=realm_assign_total_diff,

            # Profile Schema - User
            schema_df=schema_df.to_dict(orient="records"),
            schema_summary_counts=schema_summary_counts,
            schema_total_diff=schema_total_diff,

            # Profile Mappings
            mapping_df=mapping_df.to_dict(orient="records"),
            mapping_summary_counts=mapping_summary_counts,
            mapping_total_diff=mapping_total_diff,

            # Trusted Origins
            origin_df=origin_df.to_dict(orient="records"),
            origin_summary_counts=origin_summary_counts,
            origin_total_diff=origin_total_diff,

            # Event Hooks
            event_hook_df=event_hook_df.to_dict(orient="records"),
            event_hook_summary_counts=event_hook_summary_counts,
            event_hook_total_diff=event_hook_total_diff,

            # Inline Hooks
            inline_hook_df=inline_hook_df.to_dict(orient="records"),
            inline_hook_summary_counts=inline_hook_summary_counts,
            inline_hook_total_diff=inline_hook_total_diff,

            # Access Controls - Attack Protection
            attack_protection_df=attack_protection_df.to_dict(orient="records"),
            attack_protection_summary_counts=attack_protection_summary_counts,
            attack_protection_total_diff=attack_protection_total_diff,

            # Group Push Mappings
            group_push_df=group_push_df.to_dict(orient="records"),
            group_push_summary_counts=group_push_summary_counts,
            group_push_total_diff=group_push_total_diff,

            # Agents
            agent_df=agent_df.to_dict(orient="records"),
            agent_summary_counts=agent_summary_counts,
            agent_total_diff=agent_total_diff,

            # Global Session Policies (combined)
            session_df=session_df.to_dict(orient="records"),
            session_summary_counts=session_summary_counts,
            session_total_diff=session_total_diff,

            envA=envA_domain,
            envB=envB_domain,
            generated_at=datetime.now(ZoneInfo("Australia/Brisbane")).strftime(
                "%Y-%m-%d %H:%M:%S %Z"
            ),
        )


    logger.info("Rendering input form.")
    return render_template("oktacompare_form.html")


def _export_rows(rows, export_type):
    fieldnames = [
        "Entity",
        "Object",
        "Attribute",
        "Env A Value",
        "Env B Value",
        "Difference Type",
        "Impact",
        "Recommended Action",
        "Priority",
    ]
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()

    for row in rows:
        if export_type == "matches":
            value = row.get("Value", "")
            export_row = {
                "Entity": row.get("Category", ""),
                "Object": row.get("Object", ""),
                "Attribute": row.get("Attribute", ""),
                "Env A Value": value,
                "Env B Value": value,
                "Difference Type": "Match",
                "Impact": "",
                "Recommended Action": "",
                "Priority": "🟢 Match",
            }
        else:
            export_row = {
                "Entity": row.get("Category", ""),
                "Object": row.get("Object", ""),
                "Attribute": row.get("Attribute", ""),
                "Env A Value": row.get("Env A Value", ""),
                "Env B Value": row.get("Env B Value", ""),
                "Difference Type": row.get("Difference Type", ""),
                "Impact": row.get("Impact", ""),
                "Recommended Action": row.get("Recommended Action", ""),
                "Priority": row.get("Priority", ""),
            }
        writer.writerow(export_row)

    output.seek(0)
    return output


def _export_comparison_rows(diffs, matches):
    fieldnames = [
        "Category",
        "Object",
        "Attribute",
        "Env A Value",
        "Env B Value",
        "Difference Type",
        "Impact",
        "Recommended Action",
        "Priority",
    ]
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()

    def _clean_priority(value):
        if not value:
            return ""
        normalized = str(value)
        for token in ("🔴", "🟠", "🟡", "🟢"):
            normalized = normalized.replace(token, "")
        return normalized.strip()

    for row in diffs:
        writer.writerow({
            "Category": row.get("Category", ""),
            "Object": row.get("Object", ""),
            "Attribute": row.get("Attribute", ""),
            "Env A Value": row.get("Env A Value", ""),
            "Env B Value": row.get("Env B Value", ""),
            "Difference Type": row.get("Difference Type", ""),
            "Impact": row.get("Impact", ""),
            "Recommended Action": row.get("Recommended Action", ""),
            "Priority": _clean_priority(row.get("Priority", "")),
        })

    for row in matches:
        value = row.get("Value", "")
        writer.writerow({
            "Category": row.get("Category", ""),
            "Object": row.get("Object", ""),
            "Attribute": row.get("Attribute", ""),
            "Env A Value": value,
            "Env B Value": value,
            "Difference Type": "Match",
            "Impact": "",
            "Recommended Action": "",
            "Priority": "Match",
        })

    output.seek(0)
    return output


@app.route("/export_report")
def export_report():
    diffs = LAST_EXPORT.get("diffs", [])
    matches = LAST_EXPORT.get("matches", [])
    if not diffs and not matches:
        logger.warning("No comparison data found in session or server cache for export.")
    output = _export_comparison_rows(diffs, matches)
    return send_file(
        io.BytesIO(output.getvalue().encode("utf-8")),
        mimetype="text/csv",
        as_attachment=True,
        download_name="okta_compare_report.csv",
    )


@app.route("/export_differences")
def export_differences():
    diffs = LAST_EXPORT.get("diffs", [])
    if not diffs:
        logger.warning("No differences found in session or server cache for export.")
    output = _export_rows(diffs, "diffs")
    return send_file(
        io.BytesIO(output.getvalue().encode("utf-8")),
        mimetype="text/csv",
        as_attachment=True,
        download_name="okta_compare_differences.csv",
    )


@app.route("/export_matches")
def export_matches():
    matches = LAST_EXPORT.get("matches", [])
    if not matches:
        logger.warning("No matches found in session or server cache for export.")
    output = _export_rows(matches, "matches")
    return send_file(
        io.BytesIO(output.getvalue().encode("utf-8")),
        mimetype="text/csv",
        as_attachment=True,
        download_name="okta_compare_matches.csv",
    )

@app.errorhandler(requests.exceptions.ReadTimeout)
def handle_read_timeout(error):
    logger.error("Upstream timeout: %s", error)
    return render_template(
        "oktacompare_error.html",
        title="Request Timed Out",
        message="Okta took too long to respond. Please try again in a moment.",
    ), 504


@app.errorhandler(404)
def handle_not_found(error):
    logger.warning("Route not found: %s %s", request.method, request.path)
    return render_template(
        "oktacompare_error.html",
        title="Page Not Found",
        message="The requested page does not exist. Please check the URL and try again.",
    ), 404


@app.errorhandler(Exception)
def handle_unexpected_error(error):
    if isinstance(error, HTTPException):
        logger.warning("HTTP error: %s %s", error.code, error)
        return render_template(
            "oktacompare_error.html",
            title=f"Request Error ({error.code})",
            message=error.description or "The request could not be completed.",
        ), error.code
    logger.exception("Unhandled error: %s", error)
    return render_template(
        "oktacompare_error.html",
        title="Something Went Wrong",
        message="We hit an unexpected error while building the report. Please retry.",
    ), 500

@app.route("/snapshot", methods=["GET"])
def oktasnapshot_form():
    logger.info("Rendering OktaSnapshot form.")
    return render_template("oktasnapshot_form.html")


@app.route("/snapshot", methods=["POST"])
def oktasnapshot_generate():
    domain = (request.form.get("domain") or "").strip()
    api_token = (request.form.get("api_token") or "").strip()

    if not domain or not api_token:
        logger.warning("Missing required OktaSnapshot inputs.")
        return render_template(
            "oktacompare_error.html",
            title="Missing Required Input",
            message="Please provide the Okta domain and API token for OktaSnapshot.",
        ), 400

    logger.info("Generating OktaSnapshot guide for %s.", domain)
    sections, export_rows = build_oktasnapshot_guide(domain, api_token)
    OKTASNAPSHOT_EXPORT["rows"] = export_rows
    OKTASNAPSHOT_GUIDE["sections"] = sections
    OKTASNAPSHOT_GUIDE["domain"] = domain

    return redirect(url_for("oktasnapshot_guide"))


@app.route("/snapshot/guide", methods=["GET"])
def oktasnapshot_guide():
    sections = OKTASNAPSHOT_GUIDE.get("sections") or []
    domain = OKTASNAPSHOT_GUIDE.get("domain") or ""
    return render_template(
        "oktasnapshot_report.html",
        guide_sections=sections,
        guide_domain=domain,
        guide_generated_at=datetime.now(ZoneInfo("Australia/Brisbane")).strftime(
            "%Y-%m-%d %H:%M:%S %Z"
        ),
    )


@app.route("/snapshot/export", methods=["GET"])
def oktasnapshot_export():
    sections = OKTASNAPSHOT_GUIDE.get("sections") or []
    domain = OKTASNAPSHOT_GUIDE.get("domain") or ""
    if not sections:
        logger.warning("No OktaSnapshot guide data found for export.")
    try:
        from weasyprint import HTML
    except Exception:
        logger.exception("WeasyPrint not available for PDF export.")
        return render_template(
            "oktacompare_error.html",
            title="PDF Export Unavailable",
            message="PDF export requires WeasyPrint. Please install it and retry.",
        ), 500

    html = render_template(
        "oktasnapshot_pdf.html",
        guide_sections=sections,
        guide_domain=domain,
    )
    pdf = HTML(string=html).write_pdf()
    return send_file(
        io.BytesIO(pdf),
        mimetype="application/pdf",
        as_attachment=True,
        download_name="oktasnapshot_guide.pdf",
    )


def _docx_cell_value(value):
    if value is None:
        return ""
    if isinstance(value, (dict, list, tuple)):
        return json.dumps(value, sort_keys=True, indent=2, default=str)
    return str(value)


def _docx_set_cell_shading(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _docx_style_header_cell(cell, text, fill="D9E2F3"):
    from docx.shared import RGBColor

    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(31, 41, 55)
    _docx_set_cell_shading(cell, fill)


def _docx_stripe_row(table, row_idx, fill):
    for cell in table.rows[row_idx].cells:
        _docx_set_cell_shading(cell, fill)


def _docx_set_table_borders(table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9DDE5")
        borders.append(element)
    tbl_pr.append(borders)


@app.route("/snapshot/export/docx", methods=["GET"])
def oktasnapshot_export_docx():
    sections = OKTASNAPSHOT_GUIDE.get("sections") or []
    domain = OKTASNAPSHOT_GUIDE.get("domain") or ""
    if not sections:
        logger.warning("No OktaSnapshot guide data found for Word export.")
    try:
        from docx import Document
        from docx.shared import RGBColor
    except Exception:
        logger.exception("python-docx not available for Word export.")
        return render_template(
            "oktacompare_error.html",
            title="Word Export Unavailable",
            message="Word export requires python-docx. Please install it and retry.",
        ), 500

    document = Document()
    document.add_heading("OktaSnapshot Configuration", level=0)
    document.add_paragraph(f"Generated for {domain}.")

    for style_name, color in (
        ("Normal", RGBColor(0, 0, 0)),
        ("Heading 1", RGBColor(31, 41, 55)),
        ("Heading 2", RGBColor(31, 41, 55)),
    ):
        try:
            style = document.styles[style_name]
        except KeyError:
            style = None
        if style:
            style.font.color.rgb = color

    for section in sections:
        document.add_heading(section.get("title") or section.get("id") or "Section", level=1)
        rows = section.get("rows") or []
        if not rows:
            document.add_paragraph("No data available.")
            continue

        if section.get("id") in ["org-settings", "security-settings"]:
            table = document.add_table(rows=len(rows) + 1, cols=2)
            table.style = "Light Shading Accent 1"
            _docx_set_table_borders(table)
            _docx_style_header_cell(table.cell(0, 0), "Setting")
            _docx_style_header_cell(table.cell(0, 1), "Value")
            for idx, row in enumerate(rows, start=1):
                table.cell(idx, 0).text = _docx_cell_value(row.get("Setting", ""))
                table.cell(idx, 1).text = _docx_cell_value(row.get("Value", ""))
                _docx_stripe_row(table, idx, "EFF5FF" if idx % 2 == 1 else "FFFFFF")
        else:
            for row in rows:
                columns = row.keys() if section.get("id") == "applications" else section.get("columns") or row.keys()
                table = document.add_table(rows=len(columns) + 1, cols=2)
                table.style = "Light Shading Accent 1"
                _docx_set_table_borders(table)
                _docx_style_header_cell(table.cell(0, 0), "Field")
                _docx_style_header_cell(table.cell(0, 1), "Value")
                for idx, col in enumerate(columns, start=1):
                    table.cell(idx, 0).text = _docx_cell_value(col)
                    table.cell(idx, 1).text = _docx_cell_value(row.get(col, ""))
                    _docx_stripe_row(table, idx, "EFF5FF" if idx % 2 == 1 else "FFFFFF")

        document.add_paragraph("")

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return send_file(
        output,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="oktasnapshot_guide.docx",
    )


@app.route("/evaluate", methods=["GET", "POST"])
@app.route("/validate", methods=["GET", "POST"])
def okta_evaluate():
    if request.method == "POST":
        domain = (request.form.get("domain") or "").strip()
        api_token = (request.form.get("api_token") or "").strip()
        if not domain or not api_token:
            return render_template(
                "okta_evaluate.html",
                form_error="Please provide the Okta domain and API token before running evaluation.",
                form_values={"domain": domain},
            ), 400

        logger.info("Running OktaEvaluate readiness assessment for %s.", domain)
        sections, _ = build_oktasnapshot_guide(domain, api_token)
        all_apps = get_all_applications(domain, api_token, limit=200) or []
        all_users = get_users_with_security_context(domain, api_token, limit=200) or []
        api_tokens = get_api_tokens_with_metadata(domain, api_token, limit=200) or []
        custom_admin_roles = get_custom_admin_roles(domain, api_token, limit=200) or []
        resource_sets = get_resource_sets(domain, api_token, limit=200) or []
        resource_set_bindings = []
        for resource_set in resource_sets:
            resource_set_id = resource_set.get("id")
            if not resource_set_id:
                continue
            resource_set_bindings.extend(
                get_resource_set_bindings(domain, api_token, resource_set_id, limit=200) or []
            )
        result = _build_evaluate_summary(
            sections,
            domain,
            extra_context={
                "all_apps": all_apps,
                "all_users": all_users,
                "api_tokens": api_tokens,
                "custom_admin_roles": custom_admin_roles,
                "resource_set_bindings": resource_set_bindings,
            },
        )
        OKTAEVALUATE_EXPORT["evaluation"] = result
        return render_template(
            "okta_evaluate.html",
            evaluation=result,
            form_values={"domain": domain},
        )

    logger.info("Rendering OktaEvaluate page.")
    return render_template("okta_evaluate.html", form_values={})


@app.route("/evaluate/export/csv", methods=["GET"])
def okta_evaluate_export_csv():
    evaluation = OKTAEVALUATE_EXPORT.get("evaluation")
    if not evaluation:
        logger.warning("No OktaEvaluate data found for CSV export.")
    return send_file(
        io.BytesIO(_oktaevaluate_csv_bytes(evaluation or {})),
        mimetype="text/csv",
        as_attachment=True,
        download_name="oktaevaluate_security_validation_report.csv",
    )


@app.route("/evaluate/export/pdf", methods=["GET"])
def okta_evaluate_export_pdf():
    evaluation = OKTAEVALUATE_EXPORT.get("evaluation")
    if not evaluation:
        logger.warning("No OktaEvaluate data found for PDF export.")
    try:
        from weasyprint import HTML
    except Exception:
        logger.exception("WeasyPrint not available for OktaEvaluate PDF export.")
        return render_template(
            "oktacompare_error.html",
            title="PDF Export Unavailable",
            message="PDF export requires WeasyPrint. Please install it and retry.",
        ), 500

    html = render_template(
        "okta_evaluate_pdf.html",
        evaluation=evaluation or {
            "domain": "",
            "generated_at": "",
            "security_validations": [],
            "overall_score": "",
            "readiness_band": "",
            "check_prefix_legend": CHECK_PREFIX_LEGEND,
            "validation_summary": {
                "assessed": 0,
                "passed": 0,
                "high": 0,
                "medium": 0,
                "low": 0,
                "high_passed": 0,
                "medium_passed": 0,
                "low_passed": 0,
                "pass_pct": 0,
                "high_pass_pct": 0,
                "medium_pass_pct": 0,
                "low_pass_pct": 0,
            },
        },
    )
    pdf = HTML(string=html).write_pdf()
    return send_file(
        io.BytesIO(pdf),
        mimetype="application/pdf",
        as_attachment=True,
        download_name="oktaevaluate_security_validation_report.pdf",
    )


@app.route("/migrate", methods=["GET", "POST"])
def okta_migrate():
    if request.method == "POST":
        source_domain = (request.form.get("source_domain") or "").strip()
        source_token = (request.form.get("source_token") or "").strip()
        target_domain = (request.form.get("target_domain") or "").strip()
        target_token = (request.form.get("target_token") or "").strip()
        if not all([source_domain, source_token, target_domain, target_token]):
            return render_template(
                "okta_migrate.html",
                form_error="Please provide source and target domains and API tokens to generate a migration plan.",
                form_values=request.form,
            ), 400

        logger.info("Generating OktaMigrate plan for %s -> %s.", source_domain, target_domain)
        plan = _build_migration_plan(request.form)
        group_sync = None
        if request.form.get("scope_groups") == "on":
            source_groups = get_groups(source_domain, source_token) or []
            target_groups = get_groups(target_domain, target_token) or []
            group_sync = _build_group_sync_summary(source_groups, target_groups)
            OKTAMIGRATE_EXPORT["group_sync"] = group_sync
            OKTAMIGRATE_EXPORT["source_domain"] = source_domain
            OKTAMIGRATE_EXPORT["source_token"] = source_token
            OKTAMIGRATE_EXPORT["target_domain"] = target_domain
            OKTAMIGRATE_EXPORT["target_token"] = target_token
        else:
            OKTAMIGRATE_EXPORT["group_sync"] = None
            OKTAMIGRATE_EXPORT["source_domain"] = ""
            OKTAMIGRATE_EXPORT["source_token"] = ""
            OKTAMIGRATE_EXPORT["target_domain"] = ""
            OKTAMIGRATE_EXPORT["target_token"] = ""
        plan["group_sync"] = group_sync
        OKTAMIGRATE_EXPORT["plan"] = plan
        return render_template(
            "okta_migrate.html",
            migration_plan=plan,
            form_values=request.form,
        )

    logger.info("Rendering OktaMigrate page.")
    return render_template("okta_migrate.html", form_values={})


@app.route("/migrate/update/groups", methods=["POST"])
def okta_migrate_update_groups():
    plan = OKTAMIGRATE_EXPORT.get("plan")
    group_sync = OKTAMIGRATE_EXPORT.get("group_sync")
    source_domain = (OKTAMIGRATE_EXPORT.get("source_domain") or "").strip()
    source_token = (OKTAMIGRATE_EXPORT.get("source_token") or "").strip()
    target_domain = (OKTAMIGRATE_EXPORT.get("target_domain") or "").strip()
    target_token = (OKTAMIGRATE_EXPORT.get("target_token") or "").strip()
    if not plan or not group_sync or not source_domain or not source_token or not target_domain or not target_token:
        return render_template(
            "okta_migrate.html",
            form_error="No migration group comparison context found. Generate the migration plan again.",
            form_values={},
        ), 400

    selected_names = set([n for n in request.form.getlist("selected_group_names") if n])
    if not selected_names:
        plan = dict(plan)
        plan["group_sync"] = group_sync
        plan["group_update_result"] = {
            "created": [],
            "updated": [],
            "errors": ["No groups selected. Select one or more missing groups and click Migrate."],
        }
        return render_template("okta_migrate.html", migration_plan=plan, form_values={}), 400

    action_result = {"created": [], "updated": [], "errors": []}

    missing_by_name = {item.get("name"): item for item in (group_sync.get("missing", []) or [])}
    for name in selected_names:
        item = missing_by_name.get(name)
        if not item:
            action_result["errors"].append(f"Group '{name}' is not currently marked as missing in target.")
            continue
        # Missing list is already filtered to OKTA_GROUP entries via _build_group_sync_summary.
        resp = _create_group(target_domain, target_token, item.get("name"), item.get("description"))
        if resp.status_code in (200, 201):
            action_result["created"].append(item.get("name"))
        else:
            action_result["errors"].append(f"Create {item.get('name')}: {resp.status_code} {resp.text[:200]}")

    # Refresh comparison after applying updates
    source_groups = get_groups(source_domain, source_token) or []
    target_groups = get_groups(target_domain, target_token) or []
    refreshed_group_sync = _build_group_sync_summary(source_groups, target_groups)
    plan = dict(plan)
    plan["group_sync"] = refreshed_group_sync
    plan["group_update_result"] = action_result
    OKTAMIGRATE_EXPORT["plan"] = plan
    OKTAMIGRATE_EXPORT["group_sync"] = refreshed_group_sync

    return render_template(
        "okta_migrate.html",
        migration_plan=plan,
        form_values={},
    )


@app.route("/assets/<path:filename>")
def assets(filename):
    return send_from_directory("templates/static", filename)


# ---------------------------------------------------
# Run App
# ---------------------------------------------------
if __name__ == "__main__":
    logger.info("Starting OktaCompare Flask app.")
    app.run(host="0.0.0.0", debug=False, port=5000)
